import os
import json
import re
from datetime import datetime
from copy import deepcopy
from collections import defaultdict, Counter

import pdfplumber
import docx
from natasha import Segmenter, NewsEmbedding, NewsNERTagger, Doc

from html_review_report import build_html_review_report

# ============================================================
# Локальная обратимая псевдонимизация документов v3.5 / Sprint 0
# Режим по умолчанию: SaaS-safe
#
# Что изменено после теста на реальных данных:
# - телефонные RegEx стали консервативнее, чтобы не ломать номера актов/сертификатов/ИНН;
# - карты распознаются только по контексту или по группам 0000 0000 0000 0000;
# - добавлены местные телефоны вида (81371)7-80-41,доб.3215;
# - расширены признаки публичных органов: инспекция, налоговая, УФНС и т.п.;
# - усилены адреса с плохим OCR: индекс+адрес, "7 Армии ул,1да", "Ленина шд.11В";
# - добавлены дополнительные правила для ФИО: Фамилия И.О., И.О. Фамилия,
#   ФИО из 3 слов, поля "Фамилия/Имя/Отчество";
# - добавлен пользовательский файл manual_hide.txt для ручного списка фрагментов,
#   которые нужно скрывать всегда;
# - добавлен rules/manual_allow.txt для явных исключений и контекстов;
# - добавлен output/reports/report_*.md с кратким отчётом;
# - ручные правила перенесены в папку rules/ с мягкой миграцией старого manual_hide.txt.
# - Sprint 0.5: добавлен guard от ложного PASSPORT внутри служебных номеров.
# - Sprint 0.5: добавлен guard от ложного ACCOUNT внутри УИД/служебных номеров.
# - Sprint 0.6: добавлен контекстный ИНН с пробелами/дефисами и OCR-suspect ИНН с З/О.
# - Sprint 1.1: добавлен первый quality pack для SNILS, паспортов, реквизитов, адресов и дат.
# - Sprint 1.2: добавлен минимальный English profile для явных английских labels.
# ============================================================

INPUT_DIR = "input"
OUTPUT_DIR = "output"
ANON_DIR = os.path.join(OUTPUT_DIR, "anonymized")
REPORTS_DIR = os.path.join(OUTPUT_DIR, "reports")
REPORT_JSON_PATH = os.path.join(OUTPUT_DIR, "anonymization_report.json")
REPORT_DOCX_PATH = os.path.join(OUTPUT_DIR, "anonymization_report.docx")
DICT_PATH = os.path.join(OUTPUT_DIR, "project_dictionary.json")

RULES_DIR = "rules"
LEGACY_MANUAL_HIDE_PATH = "manual_hide.txt"
MANUAL_HIDE_PATH = os.path.join(RULES_DIR, "manual_hide.txt")
MANUAL_ALLOW_PATH = os.path.join(RULES_DIR, "manual_allow.txt")
PUBLIC_ORGS_PATH = os.path.join(RULES_DIR, "public_orgs.txt")
PRIVATE_ORG_MARKERS_PATH = os.path.join(RULES_DIR, "private_org_markers.txt")

FEEDBACK_DIR = "feedback"
FEEDBACK_CASES_PATH = os.path.join(FEEDBACK_DIR, "cases.jsonl")

os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(ANON_DIR, exist_ok=True)
os.makedirs(REPORTS_DIR, exist_ok=True)
os.makedirs(RULES_DIR, exist_ok=True)
os.makedirs(FEEDBACK_DIR, exist_ok=True)

print("Загрузка Natasha / NLP-моделей...")
segmenter = Segmenter()
emb = NewsEmbedding()
ner_tagger = NewsNERTagger(emb)
print("Готово.\n")

# Чем меньше число, тем выше приоритет при пересечениях
PRIORITY = {
    "PASSPORT": 10,
    "DIVISION_CODE": 11,
    "SNILS": 12,
    "INN": 9,
    "OCR_SUSPECT_INN": 9.5,
    "OGRN": 14,
    "OGRNIP": 15,
    "KPP": 16,
    "ACCOUNT": 20,
    "CARD": 21,
    "BIK": 22,
    "EMAIL": 30,
    "PHONE": 31,
    "USERNAME": 32,
    "MESSENGER_LINK": 33,
    "CADASTRAL": 40,
    "POST_INDEX": 41,
    "COORDS": 42,
    "COURT_UID": 50,
    "COURT_MATERIAL": 51,
    "COURT_CASE": 52,
    "TRACK": 53,
    "ADDRESS_DETAIL": 60,
    "DATE_BIRTH": 70,
    "DATE_DOC_ISSUE": 71,
    "DATE_REGISTRATION": 72,
    "MEDICAL_DATE": 73,
    "PERSON": 80,
    "ORG_PRIVATE": 90,
    "TEXT": 999,
}

PUBLIC_ORG_KEYWORDS = [
    "суд", "мировой судья", "прокуратура", "мвд", "умвд", "фнс", "ифнс",
    "росреестр", "роспотребнадзор", "фссп", "служба судебных приставов",
    "администрация", "правительство", "министерство", "комитет", "гас правосудие",
    "госуслуги", "полиция", "следственный комитет", "ск", "федеральная служба",
    "налоговая", "налоговой", "федеральная налоговая служба",
    "инспекция", "инспекц", "межрайонная инспекция", "межрайонной инспекции",
    "уфнс", "мфц", "роспотребнадзора", "росреестра"
]

BANK_KEYWORDS = [
    "банк", "сбербанк", "втб", "альфа-банк", "альфа банк", "тинькофф", "т-банк",
    "газпромбанк", "райффайзен", "росбанк", "совкомбанк", "почта банк"
]

PRIVATE_ORG_PREFIXES = ["ООО", "АО", "ПАО", "ЗАО", "ОАО", "АНО", "НКО", "LLC", "LTD", "INC", "JSC"]


# ------------------------------------------------------------
# Словарь проекта
# ------------------------------------------------------------

def empty_project_dictionary():
    return {
        "tokens": {},
        "reverse": {},
        "counters": {}
    }


def load_project_dictionary():
    if os.path.exists(DICT_PATH):
        with open(DICT_PATH, "r", encoding="utf-8") as f:
            data = json.load(f)
        # Поддержка нового формата
        if "tokens" in data and "reverse" in data and "counters" in data:
            return data
        # Мягкая миграция старого формата {"[TOKEN_1]": "значение"}
        migrated = empty_project_dictionary()
        for token, value in data.items():
            m = re.match(r"\[([A-Z_]+)_([0-9]+)\]", token)
            entity_type = m.group(1) if m else "TEXT"
            idx = int(m.group(2)) if m else 0
            migrated["tokens"][token] = {
                "value": value,
                "type": entity_type,
                "method": "migrated_old_dictionary"
            }
            migrated["reverse"][f"{entity_type}::{value}"] = token
            migrated["counters"][entity_type] = max(migrated["counters"].get(entity_type, 0), idx)
        return migrated
    return empty_project_dictionary()


def save_project_dictionary(project_dictionary):
    with open(DICT_PATH, "w", encoding="utf-8") as f:
        json.dump(project_dictionary, f, ensure_ascii=False, indent=4)


def get_or_create_token(value, entity_type, method, project_dictionary):
    value = value.strip()
    reverse_key = f"{entity_type}::{value}"
    if reverse_key in project_dictionary["reverse"]:
        return project_dictionary["reverse"][reverse_key]

    current = project_dictionary["counters"].get(entity_type, 0) + 1
    project_dictionary["counters"][entity_type] = current
    token = f"[{entity_type}_{current}]"

    project_dictionary["tokens"][token] = {
        "value": value,
        "type": entity_type,
        "method": method
    }
    project_dictionary["reverse"][reverse_key] = token
    return token


# ------------------------------------------------------------
# Вспомогательные функции поиска
# ------------------------------------------------------------

def make_finding(text, start, end, entity_type, method, comment=""):
    value = text[start:end]
    if not value or not value.strip():
        return None
    return {
        "value": value,
        "type": entity_type,
        "start": start,
        "end": end,
        "method": method,
        "comment": comment,
        "token": None,
    }


def add_pattern_findings(findings, text, pattern, entity_type, method, flags=0, group=0, comment=""):
    for m in re.finditer(pattern, text, flags):
        start, end = m.span(group)
        item = make_finding(text, start, end, entity_type, method, comment)
        if item:
            findings.append(item)


def remove_overlapping_findings(findings):
    """Оставляет более приоритетные и более длинные находки."""
    prepared = []
    for f in findings:
        if f["start"] >= f["end"]:
            continue
        f = deepcopy(f)
        f["priority"] = 0 if str(f.get("method", "")).startswith("manual_hide") else PRIORITY.get(f["type"], 999)
        f["length"] = f["end"] - f["start"]
        prepared.append(f)

    prepared.sort(key=lambda x: (x["priority"], -x["length"], x["start"]))
    accepted = []

    def overlaps(a, b):
        return not (a["end"] <= b["start"] or b["end"] <= a["start"])

    for candidate in prepared:
        if any(overlaps(candidate, chosen) for chosen in accepted):
            continue
        accepted.append(candidate)

    accepted.sort(key=lambda x: x["start"])
    for item in accepted:
        item.pop("priority", None)
        item.pop("length", None)
    return accepted


def apply_replacements(text, findings, project_dictionary):
    result = text
    replaced = []

    for f in sorted(findings, key=lambda x: x["start"], reverse=True):
        token = get_or_create_token(f["value"], f["type"], f["method"], project_dictionary)
        f["token"] = token
        result = result[:f["start"]] + token + result[f["end"]:]
        replaced.append(f)

    replaced.sort(key=lambda x: x["start"])
    return result, replaced



# ------------------------------------------------------------
# Пользовательский ручной список manual_hide.txt
# ------------------------------------------------------------

MANUAL_TYPE_ALIASES = {
    "PERSON": "PERSON",
    "ФИО": "PERSON",
    "ПЕРСОНА": "PERSON",
    "ЧЕЛОВЕК": "PERSON",

    "ORG": "ORG_PRIVATE",
    "ORG_PRIVATE": "ORG_PRIVATE",
    "ОРГ": "ORG_PRIVATE",
    "ОРГАНИЗАЦИЯ": "ORG_PRIVATE",
    "КОМПАНИЯ": "ORG_PRIVATE",

    "ADDRESS": "ADDRESS_DETAIL",
    "ADDRESS_DETAIL": "ADDRESS_DETAIL",
    "АДРЕС": "ADDRESS_DETAIL",

    "PHONE": "PHONE",
    "ТЕЛЕФОН": "PHONE",

    "EMAIL": "EMAIL",
    "ПОЧТА": "EMAIL",

    "INN": "INN",
    "ИНН": "INN",

    "OGRN": "OGRN",
    "ОГРН": "OGRN",

    "ACCOUNT": "ACCOUNT",
    "СЧЕТ": "ACCOUNT",
    "СЧЁТ": "ACCOUNT",

    "TEXT": "TEXT",
    "ТЕКСТ": "TEXT",
    "MANUAL": "TEXT",
    "РУЧНОЙ": "TEXT",
}


def write_text_file_if_missing(path, sample):
    """Создаёт текстовый файл с примером, если его ещё нет."""
    if os.path.exists(path):
        return
    with open(path, "w", encoding="utf-8") as f:
        f.write(sample)


def ensure_rules_files():
    """Создаёт папку rules/ и базовые файлы правил.

    Если рядом со скриптом уже есть старый manual_hide.txt, он мягко копируется
    в rules/manual_hide.txt, чтобы не потерять пользовательские ручные правила.
    """
    os.makedirs(RULES_DIR, exist_ok=True)

    manual_hide_sample = """# rules/manual_hide.txt
# Сюда можно добавить фрагменты, которые скрипт должен скрывать всегда.
# Одна строка = одно правило.
#
# Формат:
#   ТИП: значение
#
# Частые типы:
#   PERSON: Иванов Иван Иванович
#   PERSON: Иванов И.И.
#   ORG_PRIVATE: ООО Ромашка
#   ADDRESS_DETAIL: ул. Лесная, д. 5
#   PHONE: +7 921 123-45-67
#   EMAIL: example@example.com
#   TEXT: любой фрагмент, который нужно скрыть
#
# Пустые строки и строки, начинающиеся с #, игнорируются.
# Если тип не указан, будет использован TEXT.
"""

    manual_allow_sample = """# rules/manual_allow.txt
# Сюда можно добавить значения или контексты, которые НЕ нужно скрывать.
# Это файл для борьбы с ложными срабатываниями.
#
# Примеры публичных органов и служебных контекстов:
АКТ №
акт №
Сертификат:
сертификат:
Номер обращения
номер обращения
УИД
UID
регистрационный номер
ИФНС России
ФНС России
УФНС
МВД
ФССП
Росреестр
Роспотребнадзор
Тосненский городской суд
городской суд
районный суд
арбитражный суд
"""

    public_orgs_sample = """# rules/public_orgs.txt
# Публичные органы и организации, которые обычно нужно оставлять.
суд
мировой судья
городской суд
районный суд
арбитражный суд
прокуратура
МВД
УМВД
ФНС
ИФНС
УФНС
Росреестр
Роспотребнадзор
ФССП
служба судебных приставов
администрация
правительство
министерство
комитет
ГАС Правосудие
Госуслуги
полиция
Следственный комитет
СК
федеральная служба
налоговая
федеральная налоговая служба
инспекция
межрайонная инспекция
МФЦ
ЦБ РФ
Банк России
"""

    private_org_markers_sample = """# rules/private_org_markers.txt
# Маркеры частных организаций.
ООО
АО
ПАО
ЗАО
ОАО
АНО
НКО
ИП
общество с ограниченной ответственностью
акционерное общество
индивидуальный предприниматель
компания
фонд
"""

    # Мягкая миграция старого manual_hide.txt из корня проекта.
    if not os.path.exists(MANUAL_HIDE_PATH) and os.path.exists(LEGACY_MANUAL_HIDE_PATH):
        with open(LEGACY_MANUAL_HIDE_PATH, "r", encoding="utf-8") as src:
            old_content = src.read()
        with open(MANUAL_HIDE_PATH, "w", encoding="utf-8") as dst:
            dst.write(old_content)
            dst.write("\n\n# Автоматически перенесено из старого manual_hide.txt в корне проекта.\n")
    else:
        write_text_file_if_missing(MANUAL_HIDE_PATH, manual_hide_sample)

    write_text_file_if_missing(MANUAL_ALLOW_PATH, manual_allow_sample)
    write_text_file_if_missing(PUBLIC_ORGS_PATH, public_orgs_sample)
    write_text_file_if_missing(PRIVATE_ORG_MARKERS_PATH, private_org_markers_sample)


def read_rule_lines(path):
    if not os.path.exists(path):
        return []

    lines = []
    with open(path, "r", encoding="utf-8") as f:
        for raw in f:
            line = raw.strip()
            if not line or line.startswith("#"):
                continue
            lines.append(line)
    return lines


def parse_manual_hide_line(line):
    """
    Поддерживаемые варианты:
      PERSON: Иванов
      PERSON | Иванов
      PERSON = Иванов
      Иванов             -> TEXT
    """
    raw = line.strip()
    if not raw or raw.startswith("#"):
        return None

    entity_type = "TEXT"
    value = raw

    for sep in ("|", ":", "="):
        if sep in raw:
            left, right = raw.split(sep, 1)
            left_clean = left.strip().upper()
            if left_clean in MANUAL_TYPE_ALIASES:
                entity_type = MANUAL_TYPE_ALIASES[left_clean]
                value = right.strip()
            break

    if not value:
        return None

    return {
        "type": entity_type,
        "value": value,
        "raw": raw
    }


def load_manual_hide_rules():
    ensure_rules_files()
    rules = []

    for line in read_rule_lines(MANUAL_HIDE_PATH):
        rule = parse_manual_hide_line(line)
        if rule:
            rules.append(rule)

    # Более длинные фрагменты ищем первыми, чтобы manual-hide лучше перекрывал короткие совпадения.
    rules.sort(key=lambda x: len(x["value"]), reverse=True)
    return rules


def load_all_rules():
    """Загружает все пользовательские правила Sprint 0."""
    ensure_rules_files()
    rules = {
        "manual_hide": load_manual_hide_rules(),
        "manual_allow": read_rule_lines(MANUAL_ALLOW_PATH),
        "public_orgs": read_rule_lines(PUBLIC_ORGS_PATH),
        "private_org_markers": read_rule_lines(PRIVATE_ORG_MARKERS_PATH),
    }

    public_existing = {x.lower() for x in PUBLIC_ORG_KEYWORDS}
    for item in rules["public_orgs"]:
        item_clean = item.strip()
        if item_clean and item_clean.lower() not in public_existing:
            PUBLIC_ORG_KEYWORDS.append(item_clean)
            public_existing.add(item_clean.lower())

    private_existing = {x.upper() for x in PRIVATE_ORG_PREFIXES}
    for item in rules["private_org_markers"]:
        item_clean = item.strip()
        if item_clean and item_clean.upper() not in private_existing:
            PRIVATE_ORG_PREFIXES.append(item_clean)
            private_existing.add(item_clean.upper())

    return rules

def build_manual_phrase_pattern(phrase):
    """
    Точное фразовое совпадение, но пробелы внутри фразы допускают любые whitespace.
    Для слов ставим мягкие границы, чтобы 'Иван' не скрывал часть 'Иванов'.
    """
    phrase = phrase.strip()
    parts = re.split(r"\s+", phrase)
    body = r"\s+".join(re.escape(p) for p in parts if p)

    word_chars = "А-Яа-яЁёA-Za-z0-9"
    prefix = ""
    suffix = ""

    if phrase and re.match(rf"[{word_chars}]", phrase[0]):
        prefix = rf"(?<![{word_chars}])"
    if phrase and re.match(rf"[{word_chars}]", phrase[-1]):
        suffix = rf"(?![{word_chars}])"

    return prefix + body + suffix


def find_manual_hide_entities(text, manual_rules):
    findings = []
    if not manual_rules:
        return findings

    for rule in manual_rules:
        pattern = build_manual_phrase_pattern(rule["value"])
        try:
            for m in re.finditer(pattern, text, flags=re.IGNORECASE):
                item = make_finding(
                    text,
                    m.start(),
                    m.end(),
                    rule["type"],
                    "manual_hide",
                    f"manual_hide.txt: {rule['raw']}"
                )
                if item:
                    findings.append(item)
        except re.error:
            # Если пользователь случайно ввёл что-то странное, правило пропускаем.
            continue

    return findings


# ------------------------------------------------------------
# manual_allow.txt и защита от ложных срабатываний
# ------------------------------------------------------------

DEFAULT_SERVICE_NUMBER_CONTEXTS = [
    "акт №", "акт n", "акт no", "сертификат", "номер обращения",
    "обращение", "уид", "uid", "регистрационный номер", "номер дела",
    "справка", "справка №",
]

DEFAULT_NON_CARD_CONTEXTS = [
    "номер обращения", "обращение", "уид", "uid", "акт №", "акт n",
    "сертификат", "регистрационный номер"
]

DEFAULT_STRONG_PASSPORT_CONTEXTS = [
    "паспорт", "паспорта", "паспортные", "паспортный",
    "документ, удостоверяющий личность", "документ удостоверяющий личность",
    "passport", "passport no", "passport number", "passport details",
]

DEFAULT_ACCOUNT_CONTEXTS = [
    "расчётный счёт", "расчетный счет", "расчётный счет", "расчетный счёт",
    "р/с", "р\с", "расч. сч", "банковский счёт", "банковский счет",
    "корреспондентский счёт", "корреспондентский счет", "к/с", "к\с",
    "счет получателя", "счёт получателя", "счет плательщика", "счёт плательщика",
    "лицевой счет", "лицевой счёт", "банк",
    "bank account", "account number", "beneficiary account", "correspondent account"
]


def get_context(text, start, end, window=45):
    left = max(0, start - window)
    right = min(len(text), end + window)
    return text[left:right]


def normalize_for_rule_match(value):
    return re.sub(r"\s+", " ", str(value or "").lower()).strip()


def allowed_text_matches(entity_text, manual_allow):
    entity_norm = normalize_for_rule_match(entity_text)
    if not entity_norm:
        return False

    for allowed in manual_allow:
        allowed_norm = normalize_for_rule_match(allowed)
        if not allowed_norm:
            continue
        if entity_norm == allowed_norm:
            return True
        # Для длинных фраз допускаем вложенность, чтобы полное название суда можно было разрешить целиком.
        if len(allowed_norm) >= 8 and (allowed_norm in entity_norm or entity_norm in allowed_norm):
            return True
    return False


def context_contains_any(context, markers):
    context_norm = normalize_for_rule_match(context)
    return any(
        normalize_for_rule_match(marker) in context_norm
        for marker in markers
        if normalize_for_rule_match(marker)
    )


def number_context_rules(manual_allow):
    """Берёт из manual_allow только правила, похожие на контекст служебных номеров.

    Это нужно, чтобы публичные органы вроде "ИФНС России" не отключали маскировку
    телефонов или карт, если телефон/карта стоят рядом с названием органа.
    """
    markers = [
        "акт", "сертифик", "обращ", "уид", "uid", "регистрацион",
        "номер дела", "номер документа", "трек", "рпо"
    ]
    result = []
    for item in manual_allow:
        item_norm = normalize_for_rule_match(item)
        if any(marker in item_norm for marker in markers):
            result.append(item)
    return result


def is_inside_long_digit_sequence(text, finding):
    start, end = finding["start"], finding["end"]
    before = text[start - 1] if start > 0 else ""
    after = text[end] if end < len(text) else ""
    return before.isdigit() or after.isdigit()


def filter_allowed_findings(text, findings, rules):
    """Отменяет явные false positive перед применением замен.

    manual_hide не отменяется никогда.
    skipped возвращается для отчёта, чтобы было видно, какие срабатывания были подавлены.
    """
    manual_allow = rules.get("manual_allow", [])
    number_allow_contexts = number_context_rules(manual_allow)
    service_contexts = number_allow_contexts + DEFAULT_SERVICE_NUMBER_CONTEXTS
    non_card_contexts = number_allow_contexts + DEFAULT_NON_CARD_CONTEXTS

    kept = []
    skipped = []

    for finding in findings:
        method = str(finding.get("method", ""))
        entity_type = finding.get("type", "")
        context = get_context(text, finding["start"], finding["end"])

        if method.startswith("manual_hide"):
            kept.append(finding)
            continue

        if allowed_text_matches(finding.get("value", ""), manual_allow):
            item = deepcopy(finding)
            item["skip_reason"] = "manual_allow_text"
            item["context"] = context
            skipped.append(item)
            continue

        # Ложные телефоны внутри актов, сертификатов, обращений и длинных служебных номеров.
        if entity_type == "PHONE":
            if context_contains_any(context, service_contexts):
                item = deepcopy(finding)
                item["skip_reason"] = "service_number_context"
                item["context"] = context
                skipped.append(item)
                continue
            if is_inside_long_digit_sequence(text, finding):
                item = deepcopy(finding)
                item["skip_reason"] = "inside_long_digit_sequence"
                item["context"] = context
                skipped.append(item)
                continue

        # Ложные карты внутри номеров обращений/актов/сертификатов.
        if entity_type == "CARD":
            if context_contains_any(context, non_card_contexts):
                item = deepcopy(finding)
                item["skip_reason"] = "non_card_document_number_context"
                item["context"] = context
                skipped.append(item)
                continue

        # Ложные паспорта внутри служебных номеров актов/сертификатов/обращений/УИД.
        # Важно: если рядом есть явный паспортный контекст, паспортное срабатывание сохраняем.
        if entity_type == "PASSPORT":
            if context_contains_any(context, service_contexts) and not context_contains_any(context, DEFAULT_STRONG_PASSPORT_CONTEXTS):
                item = deepcopy(finding)
                item["skip_reason"] = "service_number_context"
                item["context"] = context
                skipped.append(item)
                continue

        # Ложные 20-значные счета внутри УИД, номеров обращений, актов и сертификатов.
        # Важно: если рядом есть явный банковский контекст, ACCOUNT-срабатывание сохраняем.
        if entity_type == "ACCOUNT":
            if context_contains_any(context, service_contexts) and not context_contains_any(context, DEFAULT_ACCOUNT_CONTEXTS):
                item = deepcopy(finding)
                item["skip_reason"] = "service_number_context"
                item["context"] = context
                skipped.append(item)
                continue

        # Публичные органы, добавленные в allow, не скрываем как организации.
        if entity_type in {"ORG_PRIVATE", "ORG"} and context_contains_any(context, manual_allow):
            item = deepcopy(finding)
            item["skip_reason"] = "manual_allow_org_context"
            item["context"] = context
            skipped.append(item)
            continue

        kept.append(finding)

    return kept, skipped


# ------------------------------------------------------------
# RegEx-детекторы
# ------------------------------------------------------------



OCR_INN_CHAR_MAP = {
    "З": "3",
    "з": "3",
    "О": "0",
    "о": "0",
}


def normalize_inn_candidate_digits(value):
    """Удаляет разрешённые форматные разделители внутри ИНН-кандидата."""
    return re.sub(r"[\s-]+", "", value or "")


def normalize_ocr_inn_candidate(value):
    """Минимальная OCR-нормализация только для suspect-классификации ИНН."""
    compact = normalize_inn_candidate_digits(value)
    return "".join(OCR_INN_CHAR_MAP.get(ch, ch) for ch in compact)


def add_spaced_inn_findings(findings, text):
    """Находит ИНН с пробелами/дефисами только рядом с явной меткой ИНН.

    Примеры:
      ИНН: 473 254 765 214 -> INN
      ИНН: 473-254-765-214 -> INN
      АКТ № 473 254 765 214 -> не трогаем здесь, потому что нет метки ИНН
    """
    pattern = r"\bИНН\b[\s:№-]*((?:\d[\s-]*){11}\d|(?:\d[\s-]*){9}\d)(?![\s-]*\d)"

    for m in re.finditer(pattern, text, flags=re.IGNORECASE):
        candidate = m.group(1)
        normalized = normalize_inn_candidate_digits(candidate)

        # Слитные 10/12 цифр уже обрабатывает строгий INN regex.
        if not re.search(r"[\s-]", candidate):
            continue

        if len(normalized) not in (10, 12) or not normalized.isdigit():
            continue

        item = make_finding(
            text,
            m.start(1),
            m.end(1),
            "INN",
            "regex_inn_context_spaced",
            "ИНН рядом с меткой ИНН; форматные пробелы/дефисы нормализованы"
        )
        if item:
            findings.append(item)


def add_ocr_suspect_inn_findings(findings, text):
    """Находит подозрительные OCR/форматные ИНН рядом с явной меткой ИНН.

    Важно: это НЕ расширяет валидный ИНН до произвольных длинных номеров.
    Валидные ИНН остаются только 10 или 12 цифр.

    Примеры:
      ИНН: 473254765214   -> INN
      ИНН: 4732547652143  -> OCR_SUSPECT_INN
      ИНН: 47325476521З   -> OCR_SUSPECT_INN
      АКТ № 47325476521З  -> не трогаем здесь, потому что нет метки ИНН
    """
    # 1) Старый controlled suspect: 11/13/14 цифр рядом с ИНН.
    pattern = r"\bИНН\b[\s:№-]*(\d{10,14})\b"

    for m in re.finditer(pattern, text, flags=re.IGNORECASE):
        digits = m.group(1)

        # 10 и 12 цифр обрабатывает строгий INN-детектор выше.
        # Здесь фиксируем только повреждённые/подозрительные длины рядом с меткой ИНН.
        if len(digits) in (10, 12):
            continue

        item = make_finding(
            text,
            m.start(1),
            m.end(1),
            "OCR_SUSPECT_INN",
            "regex_inn_ocr_suspect",
            "Подозрительный ИНН рядом с меткой ИНН: длина не 10/12"
        )
        if item:
            findings.append(item)

    # 2) Sprint 0.6: OCR-like character inside 10/12-position INN candidate.
    # Keep this contextual and narrow: only after explicit ИНН label.
    ocr_chars = "".join(re.escape(ch) for ch in OCR_INN_CHAR_MAP.keys())
    pattern = rf"\bИНН\b[\s:№-]*((?:[0-9{ocr_chars}][\s-]*){{11}}[0-9{ocr_chars}]|(?:[0-9{ocr_chars}][\s-]*){{9}}[0-9{ocr_chars}])(?![\s-]*[0-9])"

    for m in re.finditer(pattern, text, flags=re.IGNORECASE):
        candidate = m.group(1)
        compact = normalize_inn_candidate_digits(candidate)
        normalized = normalize_ocr_inn_candidate(candidate)

        if compact.isdigit():
            continue
        if not any(ch in OCR_INN_CHAR_MAP for ch in compact):
            continue
        if len(normalized) not in (10, 12) or not normalized.isdigit():
            continue

        item = make_finding(
            text,
            m.start(1),
            m.end(1),
            "OCR_SUSPECT_INN",
            "regex_inn_ocr_suspect_char",
            "Подозрительный ИНН рядом с меткой ИНН: OCR-похожий символ внутри значения"
        )
        if item:
            findings.append(item)


def add_contextual_spaced_digits_finding(findings, text, label_pattern, digit_count, entity_type, method, comment=""):
    """Находит длинные числовые идентификаторы с пробелами/дефисами рядом с явной меткой.

    Это controlled-fix для реквизитов вроде `БИК: 044 525 225` или
    `ОГРН: 102 770 013 2195`. Без явной метки такие числа не маскируются,
    чтобы не ломать номера актов, сертификатов и обращений.
    """
    pattern = rf"\b(?:{label_pattern})\b[\s:№#-]*((?:\d[\s-]*){{{digit_count - 1}}}\d)(?![\s-]*\d)"
    for m in re.finditer(pattern, text, flags=re.IGNORECASE):
        candidate = m.group(1)
        normalized = normalize_inn_candidate_digits(candidate)
        if len(normalized) != digit_count or not normalized.isdigit():
            continue
        item = make_finding(text, m.start(1), m.end(1), entity_type, method, comment)
        if item:
            findings.append(item)


def find_private_org_by_prefix(text):
    """Дополнительный regex для частных организаций с явным префиксом.

    Natasha иногда пропускает короткие/условные названия в кавычках:
    `ООО «Ромашка-Сервис»`. Сохраняем организационно-правовую форму, но
    скрываем название после неё.
    """
    findings = []
    prefix_pattern = r"ООО|АО|ПАО|ЗАО|ОАО|АНО|НКО|LLC|LTD|INC|JSC"
    name_token = r"[А-ЯЁA-Z0-9][А-ЯЁA-Zа-яёa-z0-9&.\-]*"
    name_pattern = rf"{name_token}(?:\s+{name_token}){{0,4}}"
    pattern = rf"\b(?:{prefix_pattern})\s+[\"«“]?({name_pattern})[\"»”]?"
    add_pattern_findings(
        findings,
        text,
        pattern,
        "ORG_PRIVATE",
        "regex_private_org_prefix",
        flags=re.IGNORECASE,
        group=1,
        comment="Частная организация по явному префиксу ООО/АО/ПАО/LLC/Ltd"
    )
    return findings

def find_regex_entities(text):
    findings = []

    # Строгие идентификаторы
    add_pattern_findings(findings, text, r"\b\d{2}\s?\d{2}\s?\d{6}\b", "PASSPORT", "regex_passport", comment="Паспорт РФ, серия/номер")
    add_pattern_findings(findings, text, r"\b\d{3}-\d{3}\b", "DIVISION_CODE", "regex_division_code", comment="Код подразделения")
    add_pattern_findings(findings, text, r"\b(?:код\s+(?:подразделения|подр\.?))[\s:№-]*(\d{3}[\s-]\d{3})\b", "DIVISION_CODE", "regex_division_code_context_spaced", flags=re.IGNORECASE, group=1, comment="Код подразделения с пробелом")
    add_pattern_findings(findings, text, r"\b\d{3}[\s-]\d{3}[\s-]\d{3}\s?\d{2}\b", "SNILS", "regex_snils", comment="СНИЛС")
    add_pattern_findings(
        findings,
        text,
        r"\b(?:паспорт|паспорта|паспортные\s+данные|документ,?\s+удостоверяющий\s+личность)\b[^\n;]{0,80}?\bсер(?:ия|ии|\.?)\s*[:№-]*((?:\d{2}\s?\d{2})[^\d\n;]{0,20}\d{6})\b",
        "PASSPORT", "regex_passport_series_number_context", flags=re.IGNORECASE, group=1,
        comment="Паспорт РФ: серия и номер в раздельных полях"
    )

    add_pattern_findings(
        findings,
        text,
        r"\b(?:passport|passport\s+no\.?|passport\s+number|passport\s+details)\b[^\n;]{0,40}?((?:\d{2}\s?\d{2})[^\d\n;]{0,15}\d{6}|\d{4}\s?\d{6})\b",
        "PASSPORT", "regex_passport_en_context", flags=re.IGNORECASE, group=1,
        comment="Passport number near English passport label"
    )

    # ИНН/ОГРН/ОГРНИП/КПП — по контексту, чтобы не ломать номера договоров и актов
    add_pattern_findings(findings, text, r"\bИНН\s*/\s*КПП\b[\s:№-]*(\d{10}|\d{12})\s*/\s*\d{9}\b", "INN", "regex_inn_kpp_pair_context", flags=re.IGNORECASE, group=1, comment="ИНН в паре ИНН/КПП")
    add_pattern_findings(findings, text, r"\bИНН\s*/\s*КПП\b[\s:№-]*(?:\d{10}|\d{12})\s*/\s*(\d{9})\b", "KPP", "regex_inn_kpp_pair_context", flags=re.IGNORECASE, group=1, comment="КПП в паре ИНН/КПП")
    add_pattern_findings(findings, text, r"\bИНН\b[\s:№-]*(\d{10}|\d{12})\b", "INN", "regex_inn_context", flags=re.IGNORECASE, group=1, comment="ИНН")
    add_pattern_findings(findings, text, r"\b(?:INN|TIN|Tax\s+ID|Taxpayer\s+ID)\b[\s:№#-]*(\d{10}|\d{12})\b", "INN", "regex_inn_en_context", flags=re.IGNORECASE, group=1, comment="INN/TIN/Tax ID near English label")
    add_spaced_inn_findings(findings, text)
    add_contextual_spaced_digits_finding(findings, text, r"INN|TIN|Tax\s+ID|Taxpayer\s+ID", 10, "INN", "regex_inn_en_context_spaced", comment="INN/TIN/Tax ID with spaces near English label")
    add_contextual_spaced_digits_finding(findings, text, r"INN|TIN|Tax\s+ID|Taxpayer\s+ID", 12, "INN", "regex_inn_en_context_spaced", comment="INN/TIN/Tax ID with spaces near English label")
    add_ocr_suspect_inn_findings(findings, text)
    add_pattern_findings(findings, text, r"\bОГРН\s*/\s*ОГРНИП\b[\s:№-]*(\d{13})\s*/\s*\d{15}\b", "OGRN", "regex_ogrn_ogrnip_pair_context", flags=re.IGNORECASE, group=1, comment="ОГРН в паре ОГРН/ОГРНИП")
    add_pattern_findings(findings, text, r"\bОГРН\s*/\s*ОГРНИП\b[\s:№-]*\d{13}\s*/\s*(\d{15})\b", "OGRNIP", "regex_ogrn_ogrnip_pair_context", flags=re.IGNORECASE, group=1, comment="ОГРНИП в паре ОГРН/ОГРНИП")
    add_pattern_findings(findings, text, r"\bОГРН\b[\s:№-]*(\d{13})\b", "OGRN", "regex_ogrn_context", flags=re.IGNORECASE, group=1, comment="ОГРН")
    add_contextual_spaced_digits_finding(findings, text, "ОГРН", 13, "OGRN", "regex_ogrn_context_spaced", comment="ОГРН с пробелами/дефисами")
    add_pattern_findings(findings, text, r"\bОГРНИП\b[\s:№-]*(\d{15})\b", "OGRNIP", "regex_ogrnip_context", flags=re.IGNORECASE, group=1, comment="ОГРНИП")
    add_contextual_spaced_digits_finding(findings, text, "ОГРНИП", 15, "OGRNIP", "regex_ogrnip_context_spaced", comment="ОГРНИП с пробелами/дефисами")
    add_pattern_findings(findings, text, r"\bКПП\b[\s:№-]*(\d{9})\b", "KPP", "regex_kpp_context", flags=re.IGNORECASE, group=1, comment="КПП")
    add_contextual_spaced_digits_finding(findings, text, "КПП", 9, "KPP", "regex_kpp_context_spaced", comment="КПП с пробелами/дефисами")

    # Банковские реквизиты
    add_pattern_findings(findings, text, r"\bБИК\b[\s:№-]*(\d{9})\b", "BIK", "regex_bik_context", flags=re.IGNORECASE, group=1, comment="БИК")
    add_contextual_spaced_digits_finding(findings, text, "БИК", 9, "BIK", "regex_bik_context_spaced", comment="БИК с пробелами/дефисами")
    add_pattern_findings(findings, text, r"\b\d{20}\b", "ACCOUNT", "regex_20_digits", comment="20-значная последовательность, скрыта как банковский счёт")

    # Банковские карты: не считаем любую слитную 16-значную последовательность картой,
    # иначе номера обращений/актов/сертификатов ошибочно превращаются в [CARD_X].
    add_pattern_findings(
        findings, text,
        r"\b(?:номер\s+карты|карта|карты|банковская\s+карта)\b[\s:№-]*(\d{13,19}|\d{4}(?:[\s-]\d{4}){3})\b",
        "CARD", "regex_card_context", flags=re.IGNORECASE, group=1, comment="Номер банковской карты по контексту"
    )
    add_pattern_findings(
        findings, text,
        r"(?<!\d)\d{4}[\s-]\d{4}[\s-]\d{4}[\s-]\d{4}(?!\d)",
        "CARD", "regex_card_grouped", comment="Номер банковской карты, записанный группами по 4"
    )

    # Контакты
    add_pattern_findings(findings, text, r"\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b", "EMAIL", "regex_email", comment="Email")

    # Телефоны: консервативно, чтобы не ломать номера актов/сертификатов/ИНН.
    add_pattern_findings(
        findings, text,
        r"(?<!\d)(?:\+7|8)[\s\-()]*\(?\d{3}\)?[\s\-()]*\d{3}[\s\-]*\d{2}[\s\-]*\d{2}(?:\s*,?\s*(?:доб\.?|доб|доп\.?)\s*\d+)?(?!\d)",
        "PHONE", "regex_phone_federal", comment="Телефон РФ +7/8"
    )
    add_pattern_findings(
        findings, text,
        r"(?<!\d)\(\d{5}\)\s*\d[\d\s\-]{4,10}(?:\s*,?\s*(?:доб\.?|доб|доп\.?)\s*\d+)?(?!\d)",
        "PHONE", "regex_phone_local_5_digit_code", comment="Местный телефон с 5-значным кодом"
    )
    add_pattern_findings(
        findings, text,
        r"\b(?:телефон|тел\.?|контактный\s+телефон|моб\.?|мобильный(?:\s+телефон)?|phone|mobile|cell(?:\s+phone)?|contact\s+phone|tel\.?)(?=[\s:№#-])[\s:№#-]*(\(?\d{3,5}\)?[\s\-]*\d[\d\s\-]{4,10}(?:\s*,?\s*(?:доб\.?|доб|доп\.?|ext\.?)\s*\d+)?)",
        "PHONE", "regex_phone_context", flags=re.IGNORECASE, group=1, comment="Телефон по контексту / phone by context"
    )

    add_pattern_findings(findings, text, r"(?<!\w)@[A-Za-z0-9_]{5,32}\b", "USERNAME", "regex_username", comment="Username / Telegram")
    add_pattern_findings(findings, text, r"\b(?:https?://)?(?:t\.me|telegram\.me|wa\.me|api\.whatsapp\.com)/[^\s,;)]+'?", "MESSENGER_LINK", "regex_messenger_link", flags=re.IGNORECASE, comment="Ссылка на мессенджер")

    # Адресные/объектные идентификаторы
    add_pattern_findings(findings, text, r"\b\d{2}:\d{2}:\d{5,7}:\d+\b", "CADASTRAL", "regex_cadastral", comment="Кадастровый номер")
    add_pattern_findings(findings, text, r"\b(?:индекс|почтовый\s+индекс)\b[\s:№-]*(\d{6})\b", "POST_INDEX", "regex_post_index_context", flags=re.IGNORECASE, group=1, comment="Почтовый индекс")
    add_pattern_findings(
        findings, text,
        r"\b(?:адрес|адрес\s+регистрации|место\s+жительства|место\s+регистрации)\b[^\n;]{0,60}?(\d{6})(?=[,\s])",
        "POST_INDEX", "regex_post_index_address_context", flags=re.IGNORECASE, group=1,
        comment="Почтовый индекс после адресного контекста"
    )
    add_pattern_findings(
        findings, text,
        r"\b(?:address|registered\s+address|residential\s+address|place\s+of\s+residence)\b[^\n;]{0,60}?(\d{6})(?=[,\s])",
        "POST_INDEX", "regex_post_index_en_address_context", flags=re.IGNORECASE, group=1,
        comment="Postal index after English address context"
    )
    # Индекс, приклеенный к адресу/OCR-тексту: 123456ТестоваяОбласть...
    add_pattern_findings(findings, text, r"(?<!\d)(\d{6})(?=\s*[,;]?\s*(?:РФ|Россия|Российская|[А-ЯЁа-яё]+ская\s+обл|[А-ЯЁа-яё]+ская\s+область|[А-ЯЁа-яё]+ская|г\.?\s*[А-ЯЁ]))", "POST_INDEX", "regex_post_index_address_attached", flags=re.IGNORECASE, group=1, comment="Почтовый индекс перед адресом")

    add_pattern_findings(findings, text, r"\b\d{1,2}\.\d{4,8},\s*\d{1,2}\.\d{4,8}\b", "COORDS", "regex_coords", comment="Координаты")

    # Судебные идентификаторы для SaaS-safe режима скрываем
    add_pattern_findings(findings, text, r"\b\d{2}RS\d{4}-\d{2}-\d{4}-\d{6}-\d{2}\b", "COURT_UID", "regex_court_uid", flags=re.IGNORECASE, comment="UID суда / ГАС")
    add_pattern_findings(findings, text, r"\b[МM]-\d+/\d{4}\b", "COURT_MATERIAL", "regex_court_material", comment="Номер судебного материала")
    add_pattern_findings(findings, text, r"\b(?:дело|гражданское\s+дело)\s*№?\s*([0-9]{1,2}-[0-9]+/[0-9]{4})\b", "COURT_CASE", "regex_court_case_context", flags=re.IGNORECASE, group=1, comment="Номер судебного дела")

    # Почтовые треки — по контексту
    add_pattern_findings(findings, text, r"\b(?:трек|трек-номер|почтов(?:ое|ый)\s+отправлени[ея]|РПО)\b[\s:№-]*(\d{14})\b", "TRACK", "regex_track_context", flags=re.IGNORECASE, group=1, comment="Почтовый трек-номер")

    return findings


# ------------------------------------------------------------
# Адресный модуль
# ------------------------------------------------------------

def find_address_details(text):
    findings = []

    # Хвост адреса от улицы/проспекта/переулка и далее.
    # ВАЖНО: более длинные маркеры стоят раньше коротких.
    street_markers = (
        r"улица|ул\.?|проспект|пр-т|пр\.?|переулок|пер\.?|шоссе|ш\.?|"
        r"набережная|наб\.?|бульвар|б-р|проезд|аллея|линия|площадь|пл\.?"
    )
    detail_markers = (
        r"квартира|кв\.?|помещение|пом\.?|корпус|корп\.?|строение|стр\.?|"
        r"дом|д\.?|офис|участок|уч\.?|к\."
    )

    # Нормальный порядок: ул. Ленина, д. 10, кв. 5
    address_tail = (
        rf"(?i)\b(?:{street_markers})\s+[^,\n;\.]+"
        rf"(?:[,;]\s*(?:{detail_markers})\s*[^,\n;\.]*)*"
    )
    add_pattern_findings(findings, text, address_tail, "ADDRESS_DETAIL", "regex_address_tail", comment="Точный адресный хвост")

    # Обратный порядок/OCR: 7 Армии ул,1да / Ленина ул., д. 10
    reverse_street_tail = (
        rf"(?i)(?:^|[,;])\s*((?:\d+\s*)?[А-ЯЁA-Zа-яёa-z0-9][А-ЯЁA-Zа-яёa-z0-9\-]*"
        rf"(?:\s+[А-ЯЁA-Zа-яёa-z0-9][А-ЯЁA-Zа-яёa-z0-9\-]*){{0,3}}\s+"
        rf"(?:{street_markers})\.?\s*,?\s*(?:(?:д\.?|дом)\s*)?\d+[А-ЯЁA-Zа-яёa-z0-9/\\-]*"
        rf"(?:[,;]\s*(?:{detail_markers})\s*[^,\n;\.]*)*)"
    )
    add_pattern_findings(findings, text, reverse_street_tail, "ADDRESS_DETAIL", "regex_reverse_street_tail", flags=re.IGNORECASE, group=1, comment="Адресный хвост: название улицы перед маркером")

    # Очень плохой OCR: Ленина шд.11В = вероятно "Ленина ш., д. 11В".
    glued_shosse_house = (
        r"(?i)(?<![А-ЯЁа-яёA-Za-z0-9])([А-ЯЁA-Zа-яёa-z0-9][А-ЯЁA-Zа-яёa-z0-9\-]{1,40}\s+ш\s*д\.?\s*\d+[А-ЯЁA-Zа-яёa-z0-9/\\-]*)"
    )
    add_pattern_findings(findings, text, glued_shosse_house, "ADDRESS_DETAIL", "regex_glued_shosse_house", flags=re.IGNORECASE, group=1, comment="Адресный хвост с плохим OCR: шд.11В")

    # Если встречаются дом/квартира/помещение без улицы в явном адресном контексте.
    detail_marker_token = rf"(?<![А-ЯЁа-яёA-Za-z])(?:{detail_markers})(?![А-ЯЁа-яёA-Za-z])"
    address_context_tail = (
        rf"(?i)\b(?:адрес|адрес\s+рег\.?|место\s+жительства|место\s+регистрации|"
        rf"место\s+нахождения\s+объекта|местонахождение\s+объекта|"
        rf"зарегистрирован(?:а)?\s+по\s+адресу)"
        rf"[^\n;]{{0,120}}?(({detail_marker_token})\s*[^,\n;\.]+"
        rf"(?:[,;]\s*{detail_marker_token}\s*[^,\n;\.]*)*)"
    )
    add_pattern_findings(findings, text, address_context_tail, "ADDRESS_DETAIL", "regex_address_context_tail", flags=re.IGNORECASE, group=1, comment="Адресные детали после явного адресного контекста")
    # Minimal English address profile: keep country/region/city context, hide street/building/apartment tail.
    english_street_markers = (
        r"street|st\.?|avenue|ave\.?|road|rd\.?|lane|ln\.?|prospect|"
        r"boulevard|blvd\.?|drive|dr\.?|square|sq\.?|embankment"
    )
    english_detail_markers = (
        r"building|bldg\.?|house|apt\.?|apartment|flat|office|suite|unit|room"
    )
    english_address_tail = (
        rf"(?i)\b(?:{english_street_markers})\s+[^,\n;\.]+"
        rf"(?:[,;]\s*(?:{english_detail_markers})\s*[^,\n;\.]*)*"
    )
    add_pattern_findings(findings, text, english_address_tail, "ADDRESS_DETAIL", "regex_address_tail_en", flags=re.IGNORECASE, comment="English address tail")

    english_reverse_street_tail = (
        rf"(?i)\b([A-Z][A-Za-z0-9.\-]*"
        rf"(?:\s+[A-Z][A-Za-z0-9.\-]*){{0,3}}\s+"
        rf"(?:{english_street_markers})"
        rf"(?:[,;]\s*(?:{english_detail_markers})\s*[^,\n;\.]*)*)"
    )
    add_pattern_findings(findings, text, english_reverse_street_tail, "ADDRESS_DETAIL", "regex_reverse_address_tail_en", flags=re.IGNORECASE, group=1, comment="English address tail: street name before marker")

    english_address_context_tail = (
        rf"(?i)\b(?:address|registered\s+address|residential\s+address|place\s+of\s+residence)"
        rf"[^\n;]{{0,120}}?((?:building|bldg\.?|house|apt\.?|apartment|flat|office|suite|unit|room)"
        rf"\s*[^,\n;\.]+(?:[,;]\s*(?:{english_detail_markers})\s*[^,\n;\.]*)*)"
    )
    add_pattern_findings(findings, text, english_address_context_tail, "ADDRESS_DETAIL", "regex_address_context_tail_en", flags=re.IGNORECASE, group=1, comment="English address detail after address context")

    return findings


# ------------------------------------------------------------
# Чувствительные даты
# ------------------------------------------------------------

def find_sensitive_dates(text):
    findings = []
    date = r"(\d{1,2}\.\d{1,2}\.\d{4})"

    add_pattern_findings(
        findings, text,
        r"(?i)(?:дата\s+рождения|дата\s+рожд\.?|родил(?:ся|ась)|рождения)\s*[:№-]?\s*" + date,
        "DATE_BIRTH", "regex_date_birth", flags=re.IGNORECASE, group=1, comment="Дата рождения"
    )
    add_pattern_findings(
        findings, text,
        r"(?i)(?:date\s+of\s+birth|birth\s+date|DOB)\s*[:№#-]?\s*" + date,
        "DATE_BIRTH", "regex_date_birth_en", flags=re.IGNORECASE, group=1, comment="Date of birth"
    )
    add_pattern_findings(
        findings, text,
        r"(?i)(?:дата\s+выдачи|паспорт\s+выдан|выдан(?:о|а)?)\D{0,80}?" + date,
        "DATE_DOC_ISSUE", "regex_date_doc_issue", flags=re.IGNORECASE, group=1, comment="Дата выдачи документа"
    )
    add_pattern_findings(
        findings, text,
        r"(?i)(?:date\s+of\s+issue|issued\s+on|passport\s+issued)\D{0,80}?" + date,
        "DATE_DOC_ISSUE", "regex_date_doc_issue_en", flags=re.IGNORECASE, group=1, comment="Document issue date"
    )
    add_pattern_findings(
        findings, text,
        r"(?i)(?:зарегистрирован(?:а)?|регистрация\s+по\s+месту\s+жительства|дата\s+регистрации)\D{0,80}?" + date,
        "DATE_REGISTRATION", "regex_date_registration", flags=re.IGNORECASE, group=1, comment="Дата регистрации"
    )
    add_pattern_findings(
        findings, text,
        r"(?i)(?:registration\s+date|registered\s+on|date\s+of\s+registration)\D{0,80}?" + date,
        "DATE_REGISTRATION", "regex_date_registration_en", flags=re.IGNORECASE, group=1, comment="Registration date"
    )
    add_pattern_findings(
        findings, text,
        r"(?i)(?:медицинск\w*|осмотр|диагноз|лечение|болезн\w*|врач|клиник\w*)\D{0,80}?" + date,
        "MEDICAL_DATE", "regex_medical_date", flags=re.IGNORECASE, group=1, comment="Медицинская дата"
    )
    return findings


# ------------------------------------------------------------
# Natasha / NLP
# ------------------------------------------------------------

def is_public_org(name):
    n = name.lower()
    return any(k in n for k in PUBLIC_ORG_KEYWORDS)


def is_bank(name):
    n = name.lower()
    return any(k in n for k in BANK_KEYWORDS)


def is_probable_english_address_location(text, start, end):
    """Skip Natasha ORG false positives for city/location words inside English address lines.

    The minimal English profile should keep city context such as
    `Registered address: 190000, Saint Petersburg, ...` while hiding
    postal index and street/building/apartment details.
    """
    value = text[start:end].strip()
    if not value:
        return False

    # Only guard Latin title-case location-like spans. Do not affect Russian ORG rules.
    if re.search(r"[?-??-???]", value):
        return False
    if not re.fullmatch(r"[A-Z][A-Za-z.\-]*(?:\s+[A-Z][A-Za-z.\-]*){0,3}", value):
        return False

    left = text[max(0, start - 120):start].lower()
    right = text[end:min(len(text), end + 140)].lower()

    has_address_label = re.search(
        r"\b(address|registered\s+address|residential\s+address|place\s+of\s+residence)\b",
        left,
    )
    if not has_address_label:
        return False

    has_address_tail = re.search(
        r"\b(street|st\.?|avenue|ave\.?|road|rd\.?|lane|ln\.?|prospect|"
        r"boulevard|blvd\.?|drive|dr\.?|square|sq\.?|embankment|"
        r"building|bldg\.?|house|apt\.?|apartment|flat|office|suite|unit|room)\b",
        right,
    )
    return bool(has_address_tail)


def split_private_org_preserve_form(name):
    """
    Возвращает (prefix, hidden_part_start_in_name, hidden_part).
    Для ООО «Ромашка» скрываем только «Ромашка», чтобы получить ООО [ORG_PRIVATE_X].
    """
    for prefix in PRIVATE_ORG_PREFIXES:
        pattern = rf"^({re.escape(prefix)})\s+(.+)$"
        m = re.match(pattern, name, flags=re.IGNORECASE)
        if m:
            return m.group(1), m.start(2), m.group(2)
    return "", 0, name


def find_role_based_persons(text):
    findings = []

    # Фамилия И.О. / И.О. Фамилия
    add_pattern_findings(
        findings, text,
        r"\b[А-ЯЁ][а-яё]+(?:-[А-ЯЁ][а-яё]+)?\s+[А-ЯЁ]\.\s?[А-ЯЁ]\.",
        "PERSON", "regex_person_surname_initials", comment="Фамилия И.О."
    )
    add_pattern_findings(
        findings, text,
        r"\b[А-ЯЁ]\.\s?[А-ЯЁ]\.\s+[А-ЯЁ][а-яё]+(?:-[А-ЯЁ][а-яё]+)?",
        "PERSON", "regex_person_initials_surname", comment="И.О. Фамилия"
    )

    # Три слова с заглавной буквы: типовой формат ФИО.
    add_pattern_findings(
        findings, text,
        r"\b[А-ЯЁ][а-яё]+(?:-[А-ЯЁ][а-яё]+)?\s+[А-ЯЁ][а-яё]+(?:-[А-ЯЁ][а-яё]+)?\s+[А-ЯЁ][а-яё]+(?:-[А-ЯЁ][а-яё]+)?\b",
        "PERSON", "regex_person_three_words", comment="Три слова с заглавной буквы, вероятное ФИО"
    )

    # ФИО/фамилия/имя после процессуальной роли или поля.
    role_or_field = (
        r"истец|ответчик|заявитель|представитель|эксперт|специалист|потребитель|"
        r"подрядчик|исполнитель|заказчик|свидетель|гражданин|гр\.?|фио|ф\.и\.о\.|"
        r"фамилия|имя|отчество"
    )
    person_value = (
        r"[А-ЯЁ][а-яё]+(?:-[А-ЯЁ][а-яё]+)?"
        r"(?:\s+[А-ЯЁ][а-яё]+(?:-[А-ЯЁ][а-яё]+)?){0,2}"
        r"|[А-ЯЁ][а-яё]+(?:-[А-ЯЁ][а-яё]+)?\s+[А-ЯЁ]\.\s?[А-ЯЁ]\."
        r"|[А-ЯЁ]\.\s?[А-ЯЁ]\.\s+[А-ЯЁ][а-яё]+(?:-[А-ЯЁ][а-яё]+)?"
    )
    add_pattern_findings(
        findings, text,
        rf"(?i)\b(?:{role_or_field})\b[\s:№-]+({person_value})",
        "PERSON", "regex_role_or_field_person", flags=re.IGNORECASE, group=1, comment="ФИО/часть ФИО после роли или поля"
    )

    add_pattern_findings(
        findings, text,
        r"(?i)\bотчество\b[\s:№-]+([А-ЯЁ][а-яё]+(?:вич|вна|ична))\b",
        "PERSON", "regex_patronymic_field", flags=re.IGNORECASE, group=1, comment="Отчество в явном поле"
    )

    english_person_value = (
        r"[A-Z][a-z]+(?:-[A-Z][a-z]+)?"
        r"(?:\s+[A-Z][a-z]+(?:-[A-Z][a-z]+)?){1,2}"
    )
    english_person_label = (
        r"full\s+name|applicant|representative|contact\s+person|"
        r"authorized\s+person|director|beneficiary"
    )
    add_pattern_findings(
        findings, text,
        rf"(?i)\b(?:{english_person_label})\b[\s:№#-]+({english_person_value})",
        "PERSON", "regex_english_labeled_person", flags=re.IGNORECASE, group=1,
        comment="English labeled person name"
    )

    return findings


def find_natasha_entities(text):
    findings = []
    doc_natasha = Doc(text)
    doc_natasha.segment(segmenter)
    doc_natasha.tag_ner(ner_tagger)

    for span in doc_natasha.spans:
        value = span.text.strip()
        if len(value) <= 2:
            continue

        if span.type == "PER":
            start = span.start
            ip_prefix = re.match(r"(?i)^ИП\s+", value)
            if ip_prefix:
                start += ip_prefix.end()
            item = make_finding(text, start, span.stop, "PERSON", "natasha_per", "ФИО / персона")
            if item:
                findings.append(item)

        elif span.type == "ORG":
            if is_probable_english_address_location(text, span.start, span.stop):
                continue
            if is_public_org(value):
                continue
            if is_bank(value):
                # По текущей политике банк оставляем, банковские реквизиты скрываем regex-слоем.
                continue

            prefix, hidden_offset, hidden_part = split_private_org_preserve_form(value)
            start = span.start + hidden_offset
            end = span.stop
            if hidden_part.strip():
                item = make_finding(text, start, end, "ORG_PRIVATE", "natasha_org_private", "Частная/непубличная организация")
                if item:
                    findings.append(item)

        # LOC не скрываем целиком: регион/город нам нужны для контекста.

    # Дополнительное правило для ИП: сохраняем ИП, скрываем ФИО после него.
    ip_pattern = r"\bИП\s+([А-ЯЁ][а-яё]+(?:\s+[А-ЯЁ][а-яё]+){1,2})\b"
    add_pattern_findings(findings, text, ip_pattern, "PERSON", "regex_ip_person", flags=re.IGNORECASE, group=1, comment="ИП: скрыто ФИО, форма ИП оставлена")

    findings.extend(find_role_based_persons(text))
    findings.extend(find_private_org_by_prefix(text))
    return findings



# ------------------------------------------------------------
# Feedback cases / human-in-the-loop review logging
# ------------------------------------------------------------

SUSPECT_ENTITY_TYPES = {"OCR_SUSPECT_INN"}
NEEDS_REVIEW_ENTITY_TYPES = set()


def finding_status_for_feedback(finding):
    entity_type = finding.get("type", "")
    if entity_type in SUSPECT_ENTITY_TYPES:
        return "SUSPECT_ENTITY"
    if entity_type in NEEDS_REVIEW_ENTITY_TYPES:
        return "NEEDS_REVIEW"
    if str(finding.get("method", "")).startswith("manual_hide"):
        return "MANUAL_CONFIRMED"
    return "VALID_ENTITY"


def policy_action_for_feedback(finding):
    status = finding_status_for_feedback(finding)
    if status == "SUSPECT_ENTITY":
        return "MASK_AND_REVIEW"
    if status == "NEEDS_REVIEW":
        return "REVIEW_ONLY"
    return "MASK_CONFIDENTLY"


def review_reason_for_feedback(finding):
    entity_type = finding.get("type", "")
    method = finding.get("method", "")

    if entity_type == "OCR_SUSPECT_INN":
        return "invalid_inn_length_near_inn_label"

    if method:
        return method

    return "review_required"


def recommended_decision_for_feedback(finding):
    status = finding_status_for_feedback(finding)
    if status == "SUSPECT_ENTITY":
        return "hide_as_suspect"
    if status == "NEEDS_REVIEW":
        return "create_candidate_rule"
    return "accept"


def human_options_for_feedback(finding):
    entity_type = finding.get("type", "")
    if entity_type == "OCR_SUSPECT_INN":
        return [
            "leave_unmasked",
            "expand_valid_inn_to_13_digits",
            "hide_as_suspect",
            "add_to_manual_allow",
            "reject_case",
        ]
    return [
        "accept",
        "hide_once",
        "hide_as_suspect",
        "add_to_manual_hide",
        "add_to_manual_allow",
        "create_candidate_rule",
        "reject_case",
    ]


def make_safe_context_for_feedback(text, finding, window=35):
    start = max(0, int(finding.get("start", 0)) - window)
    end = min(len(text), int(finding.get("end", 0)) + window)
    context = text[start:end].replace("\n", " ").replace("\r", " ")
    value = finding.get("value", "")

    digits = re.sub(r"\D", "", value)
    if digits:
        placeholder = f"<DIGITS_{len(digits)}>"
    else:
        placeholder = f"<{finding.get('type', 'VALUE')}>"

    if value:
        context = context.replace(value, placeholder, 1)

    context = re.sub(r"\s+", " ", context).strip()
    return context


def build_feedback_case(text, finding, filename, block_name):
    # Regression tests import pseudonymize.py and call anonymize_text_block directly.
    # We do not want synthetic regression runs to pollute feedback/cases.jsonl.
    if str(filename).startswith("synthetic_regression"):
        return None

    status = finding_status_for_feedback(finding)
    if status not in {"SUSPECT_ENTITY", "NEEDS_REVIEW"}:
        return None

    safe_context = make_safe_context_for_feedback(text, finding)
    review_reason = review_reason_for_feedback(finding)
    entity_type = finding.get("type", "UNKNOWN")

    import hashlib
    case_key_raw = "|".join([
        str(filename),
        str(block_name),
        str(entity_type),
        str(safe_context),
        str(review_reason),
    ])
    case_hash = hashlib.sha256(case_key_raw.encode("utf-8")).hexdigest()[:12]

    return {
        "case_id": f"case_auto_{case_hash}",
        "source": "auto_detector",
        "file_name": filename,
        "block": block_name,
        "error_type": "OCR_SUSPECT" if entity_type == "OCR_SUSPECT_INN" else "NEEDS_REVIEW",
        "entity_type_detected": entity_type,
        "entity_type_expected": entity_type,
        "status": status,
        "policy_action": policy_action_for_feedback(finding),
        "safe_context": safe_context,
        "real_value_present": True,
        "sensitivity": "high",
        "review_reason": review_reason,
        "recommended_decision": recommended_decision_for_feedback(finding),
        "human_options": human_options_for_feedback(finding),
        "created_at": datetime.now().isoformat(timespec="seconds"),
    }


def append_feedback_case(case):
    if not case:
        return False

    os.makedirs(FEEDBACK_DIR, exist_ok=True)

    existing_ids = set()
    if os.path.exists(FEEDBACK_CASES_PATH):
        try:
            with open(FEEDBACK_CASES_PATH, "r", encoding="utf-8") as f:
                for line in f:
                    line = line.strip()
                    if not line:
                        continue
                    try:
                        item = json.loads(line)
                        if item.get("case_id"):
                            existing_ids.add(item.get("case_id"))
                    except Exception:
                        continue
        except Exception:
            pass

    if case.get("case_id") in existing_ids:
        return False

    with open(FEEDBACK_CASES_PATH, "a", encoding="utf-8") as f:
        f.write(json.dumps(case, ensure_ascii=False) + "\n")

    return True

# ------------------------------------------------------------
# Главная функция анонимизации текстового блока
# ------------------------------------------------------------

def anonymize_text_block(text, filename, block_name, project_dictionary, file_report, rules):
    if not text or not text.strip():
        return text

    findings = []
    findings.extend(find_manual_hide_entities(text, rules.get("manual_hide", [])))
    findings.extend(find_regex_entities(text))
    findings.extend(find_address_details(text))
    findings.extend(find_sensitive_dates(text))
    findings.extend(find_natasha_entities(text))

    findings, skipped = filter_allowed_findings(text, findings, rules)

    for item in skipped:
        file_report["skipped"].append({
            "file": filename,
            "block": block_name,
            "type": item.get("type", ""),
            "original": item.get("value", ""),
            "method": item.get("method", ""),
            "reason": item.get("skip_reason", ""),
            "context": item.get("context", ""),
        })

    findings = remove_overlapping_findings(findings)
    anonymized_text, replaced = apply_replacements(text, findings, project_dictionary)

    for f in replaced:
        status = finding_status_for_feedback(f)
        policy_action = policy_action_for_feedback(f)
        review_reason = review_reason_for_feedback(f)

        file_report["replacements"].append({
            "file": filename,
            "block": block_name,
            "type": f["type"],
            "original": f["value"],
            "token": f["token"],
            "method": f["method"],
            "comment": f.get("comment", ""),
            "status": status,
            "policy_action": policy_action,
            "review_reason": review_reason,
        })
        file_report["summary"][f["type"]] += 1

        feedback_case = build_feedback_case(text, f, filename, block_name)
        if feedback_case:
            file_report.setdefault("review_cases", []).append(feedback_case)
            append_feedback_case(feedback_case)

    return anonymized_text


# ------------------------------------------------------------
# DOCX/PDF обработка
# ------------------------------------------------------------


TABLE_CONTEXT_LABEL_PATTERN = re.compile(
    r"(?i)(?:"
    r"\bфио\b|ф\.\s*и\.\s*о\.|\bфамилия\b|\bимя\b|\bотчество\b|"
    r"\bистец\b|\bответчик\b|\bзаявитель\b|\bпредставитель\b|\bпотребитель\b|\bгражданин\b|\bгр\.\b|"
    r"\bинн\b|\bогрн\b|\bогрнип\b|\bкпп\b|\bснилс\b|"
    r"\bпаспорт\b|паспортн|серия|номер\s+паспорта|код\s+подразделения|"
    r"телефон|тел\.|моб\.|email|e-mail|почта|"
    r"адрес|место\s+жительства|место\s+регистрации|регистрац|"
    r"дата\s+рождения|рождени|дата\s+выдачи|выдан|"
    r"банк|бик|расч[её]тн|р/с|к/с|сч[её]т|карта|"
    r"кадастров|индекс|трек|рпо|uid|уид|номер\s+дела|дело"
    r")"
)


def normalize_table_context_label(value):
    """Возвращает короткую безопасную метку строки/колонки для DOCX-таблиц.

    Важно: не каждая первая ячейка строки является заголовком. В обычной
    таблице первая колонка может содержать ФИО. Поэтому в контекст берём только
    короткие label-like значения с известными маркерами: `ФИО`, `ИНН`,
    `Паспорт`, `Адрес`, `Телефон` и т.п.
    """
    value = re.sub(r"\s+", " ", str(value or "")).strip(" \t\r\n:;№-–—")
    if not value:
        return ""
    # Длинная ячейка почти наверняка не заголовок, а обычный текст.
    if len(value) > 80:
        return ""
    if not TABLE_CONTEXT_LABEL_PATTERN.search(value):
        return ""
    return value


def unique_context_labels(labels, cell_text=""):
    """Удаляет пустые и дублирующиеся метки контекста для ячейки таблицы."""
    result = []
    seen = set()
    cell_norm = normalize_for_rule_match(cell_text)

    for label in labels:
        label = normalize_table_context_label(label)
        if not label:
            continue
        label_norm = normalize_for_rule_match(label)
        if not label_norm or label_norm == cell_norm or label_norm in seen:
            continue
        seen.add(label_norm)
        result.append(label)

    return result


def anonymize_table_cell_text(cell_text, context_labels, filename, block_name, project_dictionary, file_report, rules):
    """Анонимизирует значение ячейки с учётом заголовка строки/колонки.

    Пример: в таблице значение `473254765214` под заголовком `ИНН`
    временно анализируется как `ИНН: 473254765214`, но в DOCX возвращается
    только значение ячейки: `[INN_1]`. Это закрывает false negative, когда
    контекст находится в соседней ячейке таблицы.
    """
    labels = unique_context_labels(context_labels, cell_text)
    if not labels:
        return anonymize_text_block(cell_text, filename, block_name, project_dictionary, file_report, rules)

    prefix = " / ".join(labels) + ": "
    contextual_text = prefix + cell_text
    anonymized = anonymize_text_block(contextual_text, filename, block_name, project_dictionary, file_report, rules)

    if anonymized.startswith(prefix):
        return anonymized[len(prefix):]

    # Защитный fallback: заголовки таблиц обычно не маскируются, но если формат
    # неожиданно изменился, не вставляем служебный префикс обратно в документ.
    if ": " in anonymized:
        return anonymized.split(": ", 1)[1]

    file_report["warnings"].append(
        f"Не удалось безопасно удалить табличный контекст из блока {block_name}. Ячейка обработана без контекста."
    )
    return anonymize_text_block(cell_text, filename, block_name, project_dictionary, file_report, rules)


def process_docx(filepath, filename, project_dictionary, rules):
    document = docx.Document(filepath)
    file_report = {
        "summary": defaultdict(int),
        "warnings": [],
        "replacements": [],
        "skipped": [],
        "review_cases": []
    }

    # Параграфы основного текста
    for i, para in enumerate(document.paragraphs):
        if para.text:
            para.text = anonymize_text_block(para.text, filename, f"paragraph_{i+1}", project_dictionary, file_report, rules)

    # Таблицы
    # В DOCX табличный контекст часто находится не в самой ячейке, а в заголовке
    # колонки или строки: например, значение `473254765214` стоит под колонкой
    # `ИНН`. Поэтому для ячеек данных временно добавляем короткий контекст
    # заголовка при анализе, но сохраняем в DOCX только исходное значение ячейки
    # с заменами.
    for ti, table in enumerate(document.tables):
        rows = list(table.rows)
        column_headers = []
        if rows:
            column_headers = [normalize_table_context_label(cell.text) for cell in rows[0].cells]

        for ri, row in enumerate(rows):
            row_header = ""
            if row.cells:
                row_header = normalize_table_context_label(row.cells[0].text)

            for ci, cell in enumerate(row.cells):
                if not cell.text:
                    continue

                context_labels = []
                if ri > 0 and ci < len(column_headers):
                    context_labels.append(column_headers[ci])
                if ci > 0:
                    context_labels.append(row_header)

                block_name = f"table_{ti+1}_r{ri+1}_c{ci+1}"
                cell.text = anonymize_table_cell_text(
                    cell.text,
                    context_labels,
                    filename,
                    block_name,
                    project_dictionary,
                    file_report,
                    rules,
                )

    # Колонтитулы — обрабатываем по возможности
    for si, section in enumerate(document.sections):
        for hi, para in enumerate(section.header.paragraphs):
            if para.text:
                para.text = anonymize_text_block(para.text, filename, f"header_s{si+1}_p{hi+1}", project_dictionary, file_report, rules)
        for fi, para in enumerate(section.footer.paragraphs):
            if para.text:
                para.text = anonymize_text_block(para.text, filename, f"footer_s{si+1}_p{fi+1}", project_dictionary, file_report, rules)

    if len(document.inline_shapes) > 0:
        file_report["warnings"].append(f"В документе есть изображения: {len(document.inline_shapes)}. Текст внутри изображений не проверялся.")

    out_name = filename.rsplit(".", 1)[0] + "_anonymized.docx"
    out_path = os.path.join(ANON_DIR, out_name)
    document.save(out_path)
    return file_report, out_path


def process_pdf(filepath, filename, project_dictionary, rules):
    file_report = {
        "summary": defaultdict(int),
        "status": "pending",
        "warnings": [
            "PDF преобразуется в DOCX простым текстом. Сложная верстка PDF не сохраняется.",
            "PDF на v0.1 обрабатывается только по извлекаемому текстовому слою. OCR не выполняется."
        ],
        "pdf_extraction": {
            "pages_total": 0,
            "pages_with_text": 0,
            "pages_without_text": [],
            "chars_by_page": [],
            "total_extracted_chars": 0,
        },
        "replacements": [],
        "skipped": [],
        "review_cases": []
    }

    lines = []
    extraction = file_report["pdf_extraction"]

    with pdfplumber.open(filepath) as pdf:
        extraction["pages_total"] = len(pdf.pages)
        for page_num, page in enumerate(pdf.pages, start=1):
            extracted = page.extract_text() or ""
            stripped = extracted.strip()
            chars = len(stripped)

            extraction["chars_by_page"].append({
                "page": page_num,
                "chars": chars,
            })
            extraction["total_extracted_chars"] += chars

            if stripped:
                extraction["pages_with_text"] += 1
                anonymized = anonymize_text_block(extracted, filename, f"pdf_page_{page_num}", project_dictionary, file_report, rules)
                lines.append(anonymized)
            else:
                extraction["pages_without_text"].append(page_num)

    if extraction["total_extracted_chars"] == 0:
        file_report["status"] = "not_processed_no_text_layer"
        file_report["warnings"].append(
            "В PDF не найден извлекаемый текстовый слой. Вероятно, это скан, изображение или PDF без text layer. "
            "Anonymized DOCX не создан, чтобы не создавать ложное ощущение успешной обработки."
        )
        return file_report, None

    if extraction["pages_without_text"]:
        file_report["status"] = "partially_processed_text_layer"
        file_report["warnings"].append(
            "Часть страниц PDF не содержит извлекаемого текста и не была проверена: "
            + ", ".join(str(x) for x in extraction["pages_without_text"])
        )
    else:
        file_report["status"] = "processed_text_layer"

    out_doc = docx.Document()
    for block in lines:
        for line in block.split("\n"):
            out_doc.add_paragraph(line)

    out_name = filename.rsplit(".", 1)[0] + "_anonymized.docx"
    out_path = os.path.join(ANON_DIR, out_name)
    out_doc.save(out_path)
    return file_report, out_path


# ------------------------------------------------------------
# Отчёты
# ------------------------------------------------------------

def normalize_report_for_json(report):
    normalized = {"files": {}}
    for filename, data in report["files"].items():
        normalized["files"][filename] = {
            "summary": dict(data["summary"]),
            "status": data.get("status", "processed"),
            "warnings": data.get("warnings", []),
            "pdf_extraction": data.get("pdf_extraction"),
            "replacements": data.get("replacements", []),
            "skipped": data.get("skipped", []),
            "review_cases": data.get("review_cases", [])
        }
    return normalized


def save_json_report(report):
    with open(REPORT_JSON_PATH, "w", encoding="utf-8") as f:
        json.dump(normalize_report_for_json(report), f, ensure_ascii=False, indent=4)


def save_markdown_report(report):
    os.makedirs(REPORTS_DIR, exist_ok=True)
    report_name = f"report_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}.md"
    report_path = os.path.join(REPORTS_DIR, report_name)

    lines = []
    lines.append("# Отчёт псевдонимизации")
    lines.append("")
    lines.append(f"Дата обработки: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append("Версия: v3.5 / Sprint 0")
    lines.append("")
    lines.append("> Внимание: отчёт может содержать исходные персональные/конфиденциальные данные. Не отправляйте его в SaaS и третьим лицам.")
    lines.append("")

    total_summary = Counter()
    total_methods = Counter()
    total_skipped = Counter()

    for filename, data in report["files"].items():
        for entity_type, count in data.get("summary", {}).items():
            total_summary[entity_type] += count
        for r in data.get("replacements", []):
            total_methods[r.get("method", "unknown")] += 1
        for r in data.get("skipped", []):
            total_skipped[r.get("reason", "unknown")] += 1

    lines.append("## 1. Общая сводка замен")
    lines.append("")
    if total_summary:
        lines.append("| Тип | Количество |")
        lines.append("|---|---:|")
        for key, count in total_summary.most_common():
            lines.append(f"| {key} | {count} |")
    else:
        lines.append("Замен не найдено.")
    lines.append("")

    total_review_cases = []
    for filename, data in report["files"].items():
        total_review_cases.extend(data.get("review_cases", []))

    lines.append("## 2. Скрыто как подозрительное / требует review")
    lines.append("")
    if total_review_cases:
        lines.append("| Файл | Блок | Тип | Safe context | Причина |")
        lines.append("|---|---|---|---|---|")
        for c in total_review_cases[:50]:
            lines.append(f"| {c.get('file_name', '')} | {c.get('block', '')} | {c.get('entity_type_expected', '')} | `{c.get('safe_context', '')}` | {c.get('review_reason', '')} |")
        if len(total_review_cases) > 50:
            lines.append(f"")
            lines.append(f"Показано 50 из {len(total_review_cases)} review-кейсов. Полный список см. в feedback/cases.jsonl и anonymization_report.json.")
    else:
        lines.append("Подозрительных/review-кейсов нет.")
    lines.append("")

    lines.append("## 3. Источники срабатываний")
    lines.append("")
    if total_methods:
        lines.append("| Метод | Количество |")
        lines.append("|---|---:|")
        for key, count in total_methods.most_common():
            lines.append(f"| {key} | {count} |")
    else:
        lines.append("Нет данных.")
    lines.append("")

    lines.append("## 3. Отменённые срабатывания")
    lines.append("")
    if total_skipped:
        lines.append("| Причина | Количество |")
        lines.append("|---|---:|")
        for key, count in total_skipped.most_common():
            lines.append(f"| {key} | {count} |")
    else:
        lines.append("Отменённых срабатываний нет.")
    lines.append("")

    lines.append("## 4. Файлы")
    lines.append("")
    for filename, data in report["files"].items():
        lines.append(f"### {filename}")
        lines.append("")
        if data.get("status"):
            lines.append(f"Статус обработки: `{data.get('status')}`")
            lines.append("")

        pdf_extraction = data.get("pdf_extraction")
        if pdf_extraction:
            lines.append("PDF text-layer extraction:")
            lines.append(f"- страниц всего: {pdf_extraction.get('pages_total', 0)}")
            lines.append(f"- страниц с текстом: {pdf_extraction.get('pages_with_text', 0)}")
            lines.append(f"- извлечено символов: {pdf_extraction.get('total_extracted_chars', 0)}")
            if pdf_extraction.get("pages_without_text"):
                pages = ", ".join(str(x) for x in pdf_extraction.get("pages_without_text", []))
                lines.append(f"- страницы без text layer: {pages}")
            lines.append("")

        if data.get("summary"):
            lines.append("Замены:")
            for entity_type, count in sorted(data["summary"].items()):
                lines.append(f"- {entity_type}: {count}")
        else:
            lines.append("Замен не найдено.")

        if data.get("skipped"):
            lines.append("")
            lines.append("Отменённые срабатывания:")
            for item in data.get("skipped", [])[:30]:
                lines.append(f"- {item.get('type', '')}: `{item.get('original', '')}` — {item.get('reason', '')} ({item.get('block', '')})")

        if data.get("review_cases"):
            lines.append("")
            lines.append("Скрыто как подозрительное / требует review:")
            for item in data.get("review_cases", [])[:30]:
                lines.append(f"- {item.get('entity_type_expected', '')}: `{item.get('safe_context', '')}` — {item.get('review_reason', '')} ({item.get('block', '')})")

        if data.get("warnings"):
            lines.append("")
            lines.append("Предупреждения:")
            for warning in data["warnings"]:
                lines.append(f"- {warning}")
        lines.append("")

    lines.append("## 5. Общие предупреждения")
    lines.append("")
    lines.append("- project_dictionary.json позволяет восстановить исходные данные и должен храниться как секретный файл.")
    lines.append("- rules/manual_hide.txt может содержать реальные персональные данные.")
    lines.append("- rules/manual_allow.txt может содержать чувствительные контексты дела.")
    lines.append("- Инструмент выполняет псевдонимизацию / masking / risk reduction, а не гарантированную анонимизацию.")
    lines.append("- PDF на v0.1 обрабатывается только по текстовому слою; изображения и сканы не проверяются.")
    lines.append("- Перед отправкой результата во внешний AI/SaaS нужна ручная проверка.")
    lines.append("")

    with open(report_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

    return report_path


def save_html_report(report, project_dictionary):
    os.makedirs(REPORTS_DIR, exist_ok=True)
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    report_name = f"review_report_{timestamp}.html"
    report_path = os.path.join(REPORTS_DIR, report_name)
    latest_path = os.path.join(REPORTS_DIR, "review_report_latest.html")

    generated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    html = build_html_review_report(
        normalize_report_for_json(report),
        project_dictionary=project_dictionary,
        generated_at=generated_at,
    )

    with open(report_path, "w", encoding="utf-8") as f:
        f.write(html)
    with open(latest_path, "w", encoding="utf-8") as f:
        f.write(html)

    return report_path


def save_docx_report(report):
    report_doc = docx.Document()
    report_doc.add_heading("Отчёт анонимизации", level=1)
    report_doc.add_paragraph("Внимание: этот отчёт содержит исходные персональные/конфиденциальные данные. Не отправляйте его в SaaS и третьим лицам.")

    for filename, data in report["files"].items():
        report_doc.add_heading(f"Файл: {filename}", level=2)

        if data.get("status"):
            report_doc.add_paragraph(f"Статус обработки: {data.get('status')}")

        pdf_extraction = data.get("pdf_extraction")
        if pdf_extraction:
            report_doc.add_paragraph(
                "PDF text-layer extraction: "
                f"страниц всего {pdf_extraction.get('pages_total', 0)}, "
                f"страниц с текстом {pdf_extraction.get('pages_with_text', 0)}, "
                f"извлечено символов {pdf_extraction.get('total_extracted_chars', 0)}."
            )
            if pdf_extraction.get("pages_without_text"):
                pages = ", ".join(str(x) for x in pdf_extraction.get("pages_without_text", []))
                report_doc.add_paragraph(f"Страницы без text layer: {pages}")

        report_doc.add_heading("1. Сводка", level=3)
        if data["summary"]:
            for entity_type, count in sorted(data["summary"].items()):
                report_doc.add_paragraph(f"{entity_type}: {count}")
        else:
            report_doc.add_paragraph("Замен не найдено.")

        report_doc.add_heading("2. Таблица замен", level=3)
        table = report_doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        headers = ["Тип", "Исходное значение", "Токен", "Метод", "Блок", "Комментарий"]
        for idx, h in enumerate(headers):
            table.rows[0].cells[idx].text = h

        for r in data["replacements"]:
            cells = table.add_row().cells
            cells[0].text = r.get("type", "")
            cells[1].text = r.get("original", "")
            cells[2].text = r.get("token", "")
            cells[3].text = r.get("method", "")
            cells[4].text = r.get("block", "")
            cells[5].text = r.get("comment", "")

        report_doc.add_heading("3. Скрыто как подозрительное / требует review", level=3)
        review_cases = data.get("review_cases", [])
        if review_cases:
            table = report_doc.add_table(rows=1, cols=5)
            table.style = "Table Grid"
            headers = ["Тип", "Safe context", "Причина", "Блок", "Рекомендация"]
            for idx, h in enumerate(headers):
                table.rows[0].cells[idx].text = h
            for c in review_cases:
                cells = table.add_row().cells
                cells[0].text = c.get("entity_type_expected", "")
                cells[1].text = c.get("safe_context", "")
                cells[2].text = c.get("review_reason", "")
                cells[3].text = c.get("block", "")
                cells[4].text = c.get("recommended_decision", "")
        else:
            report_doc.add_paragraph("Подозрительных/review-кейсов нет.")

        report_doc.add_heading("4. Отменённые срабатывания manual_allow / safety filters", level=3)
        skipped = data.get("skipped", [])
        if skipped:
            table = report_doc.add_table(rows=1, cols=5)
            table.style = "Table Grid"
            headers = ["Тип", "Значение", "Метод", "Причина", "Блок"]
            for idx, h in enumerate(headers):
                table.rows[0].cells[idx].text = h
            for r in skipped:
                cells = table.add_row().cells
                cells[0].text = r.get("type", "")
                cells[1].text = r.get("original", "")
                cells[2].text = r.get("method", "")
                cells[3].text = r.get("reason", "")
                cells[4].text = r.get("block", "")
        else:
            report_doc.add_paragraph("Отменённых срабатываний нет.")

        report_doc.add_heading("5. Предупреждения", level=3)
        if data["warnings"]:
            for warning in data["warnings"]:
                report_doc.add_paragraph(warning, style=None)
        else:
            report_doc.add_paragraph("Нет предупреждений.")

        report_doc.add_heading("6. Что проверить вручную", level=3)
        manual_checks = [
            "Адреса, написанные без слов 'улица', 'дом', 'квартира'.",
            "Фамилии без имени/инициалов, если Natasha и дополнительные RegEx их не распознали.",
            "Имена/отчества без контекста: их автоматическое скрытие может давать ложные срабатывания.",
            "Изображения и сканы: OCR изображений не выполнялся.",
            "Редкие форматы банковских реквизитов и номеров документов.",
            "Судебные номера: проверьте, все ли идентификаторы дела скрыты для SaaS-режима."
        ]
        for item in manual_checks:
            report_doc.add_paragraph(item)

    report_doc.save(REPORT_DOCX_PATH)


# ------------------------------------------------------------
# Запуск
# ------------------------------------------------------------

def main():
    project_dictionary = load_project_dictionary()
    rules = load_all_rules()
    print(f"Правил rules/manual_hide.txt загружено: {len(rules.get('manual_hide', []))}")
    print(f"Правил rules/manual_allow.txt загружено: {len(rules.get('manual_allow', []))}")
    full_report = {"files": {}}

    if not os.path.exists(INPUT_DIR):
        print(f"Папка {INPUT_DIR} не найдена. Создайте её и положите туда PDF/DOCX.")
        return

    files = [f for f in os.listdir(INPUT_DIR) if f.lower().endswith((".pdf", ".docx")) and not f.startswith("~$")]
    if not files:
        print(f"В папке {INPUT_DIR} нет PDF/DOCX файлов.")
        return

    for filename in files:
        filepath = os.path.join(INPUT_DIR, filename)
        print(f"Обработка файла: {filename}")

        try:
            if filename.lower().endswith(".docx"):
                file_report, out_path = process_docx(filepath, filename, project_dictionary, rules)
            else:
                file_report, out_path = process_pdf(filepath, filename, project_dictionary, rules)

            full_report["files"][filename] = file_report
            if out_path:
                print(f"  Готово: {out_path}")
            else:
                print(f"  DOCX не создан: {file_report.get('status', 'not_processed')}")
        except Exception as e:
            full_report["files"][filename] = {
                "summary": defaultdict(int),
                "status": "processing_error",
                "warnings": [f"Ошибка обработки файла: {e}"],
                "replacements": [],
                "skipped": [],
                "review_cases": []
            }
            print(f"  ОШИБКА при обработке {filename}: {e}")

    save_project_dictionary(project_dictionary)
    save_json_report(full_report)
    save_docx_report(full_report)
    markdown_report_path = save_markdown_report(full_report)
    html_report_path = save_html_report(full_report, project_dictionary)

    print("\nАнонимизация завершена.")
    print(f"Анонимизированные файлы: {ANON_DIR}")
    print(f"Словарь проекта: {DICT_PATH}")
    print(f"Отчёт DOCX: {REPORT_DOCX_PATH}")
    print(f"Отчёт JSON: {REPORT_JSON_PATH}")
    print(f"Краткий отчёт MD: {markdown_report_path}")
    print(f"HTML review report: {html_report_path}")
    print(f"Ручной список скрытия: {MANUAL_HIDE_PATH}")
    print(f"Ручной список разрешений: {MANUAL_ALLOW_PATH}")
    print("\nВАЖНО: project_dictionary.json, anonymization_report.*, output/reports/*.md, output/reports/*.html и rules/*.txt не отправлять в SaaS.")


if __name__ == "__main__":
    main()
