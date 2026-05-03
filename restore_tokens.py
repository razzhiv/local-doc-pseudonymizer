import os
import json
import re
import docx

# ============================================================
# Восстановление токенов в DOCX по project_dictionary.json v3.1
# ============================================================

TO_DECODE_DIR = "to_decode"
OUTPUT_DIR = "output"
RESTORED_DIR = os.path.join(OUTPUT_DIR, "restored")
DICT_PATH = os.path.join(OUTPUT_DIR, "project_dictionary.json")

os.makedirs(RESTORED_DIR, exist_ok=True)


def load_project_dictionary():
    if not os.path.exists(DICT_PATH):
        raise FileNotFoundError(f"Не найден словарь: {DICT_PATH}")

    with open(DICT_PATH, "r", encoding="utf-8") as f:
        data = json.load(f)

    # Новый формат
    if "tokens" in data:
        return {token: meta["value"] for token, meta in data["tokens"].items()}

    # Старый формат {"[TOKEN_1]": "значение"}
    return data


def restore_text(text, token_map):
    if not text:
        return text

    # Сначала длинные токены, чтобы не было конфликтов.
    for token in sorted(token_map.keys(), key=len, reverse=True):
        text = text.replace(token, token_map[token])
    return text


def restore_docx(filepath, filename, token_map):
    document = docx.Document(filepath)

    for para in document.paragraphs:
        if para.text:
            para.text = restore_text(para.text, token_map)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    cell.text = restore_text(cell.text, token_map)

    for section in document.sections:
        for para in section.header.paragraphs:
            if para.text:
                para.text = restore_text(para.text, token_map)
        for para in section.footer.paragraphs:
            if para.text:
                para.text = restore_text(para.text, token_map)

    out_name = filename.rsplit(".", 1)[0] + "_restored.docx"
    out_path = os.path.join(RESTORED_DIR, out_name)
    document.save(out_path)
    return out_path


def main():
    token_map = load_project_dictionary()

    if not os.path.exists(TO_DECODE_DIR):
        print(f"Папка {TO_DECODE_DIR} не найдена. Создайте её и положите туда DOCX для восстановления.")
        return

    files = [f for f in os.listdir(TO_DECODE_DIR) if f.lower().endswith(".docx") and not f.startswith("~$")]
    if not files:
        print(f"В папке {TO_DECODE_DIR} нет DOCX файлов для восстановления.")
        return

    for filename in files:
        filepath = os.path.join(TO_DECODE_DIR, filename)
        print(f"Восстановление файла: {filename}")
        try:
            out_path = restore_docx(filepath, filename, token_map)
            print(f"  Готово: {out_path}")
        except Exception as e:
            print(f"  ОШИБКА при восстановлении {filename}: {e}")

    print("\nВосстановление токенов завершено.")
    print(f"Восстановленные файлы: {RESTORED_DIR}")


if __name__ == "__main__":
    main()
