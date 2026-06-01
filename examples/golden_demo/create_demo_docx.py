from __future__ import annotations

import sys
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[2]
FIXTURE_PATH = Path(__file__).with_name("source_synthetic_client_matter.txt")
DEFAULT_OUTPUT_PATH = ROOT / "input" / "golden_synthetic_client_matter.docx"


def _flush_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return

    width = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=width)
    for row_index, row in enumerate(rows):
        for col_index in range(width):
            table.cell(row_index, col_index).text = row[col_index] if col_index < len(row) else ""


def build_demo_docx(destination: str | Path = DEFAULT_OUTPUT_PATH) -> Path:
    """Create a disposable DOCX demo input from the committed synthetic text fixture."""
    destination = Path(destination)
    destination.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    pending_table: list[list[str]] = []

    for raw_line in FIXTURE_PATH.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if line.startswith("|") and line.endswith("|"):
            cells = [cell.strip() for cell in line.strip("|").split("|")]
            pending_table.append(cells)
            continue

        _flush_table(document, pending_table)
        pending_table = []

        if not line:
            continue
        if not document.paragraphs:
            document.add_heading(line, level=1)
        elif line.endswith(":"):
            document.add_heading(line[:-1], level=2)
        else:
            document.add_paragraph(line)

    _flush_table(document, pending_table)
    document.save(destination)
    return destination


def main(argv: list[str]) -> int:
    destination = Path(argv[1]) if len(argv) > 1 else DEFAULT_OUTPUT_PATH
    out_path = build_demo_docx(destination)
    print(f"Created synthetic demo DOCX: {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
