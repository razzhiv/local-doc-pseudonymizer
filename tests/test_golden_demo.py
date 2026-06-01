from __future__ import annotations

from pathlib import Path

from docx import Document

import pseudonymize
from examples.golden_demo.create_demo_docx import build_demo_docx


ROOT = Path(__file__).resolve().parents[1]
EXPECTED_PREVIEW_PATH = ROOT / "examples" / "golden_demo" / "expected_masked_preview.txt"


def basic_rules() -> dict:
    return {
        "manual_hide": [],
        "manual_allow": [],
        "public_orgs": [],
        "private_org_markers": [],
    }


def configure_isolated_runtime(tmp_path, monkeypatch) -> Path:
    output_dir = tmp_path / "output"
    anon_dir = output_dir / "anonymized"
    reports_dir = output_dir / "reports"
    feedback_dir = tmp_path / "feedback"

    anon_dir.mkdir(parents=True, exist_ok=True)
    reports_dir.mkdir(parents=True, exist_ok=True)
    feedback_dir.mkdir(parents=True, exist_ok=True)

    monkeypatch.setattr(pseudonymize, "OUTPUT_DIR", str(output_dir))
    monkeypatch.setattr(pseudonymize, "ANON_DIR", str(anon_dir))
    monkeypatch.setattr(pseudonymize, "REPORTS_DIR", str(reports_dir))
    monkeypatch.setattr(pseudonymize, "FEEDBACK_DIR", str(feedback_dir))
    monkeypatch.setattr(pseudonymize, "FEEDBACK_CASES_PATH", str(feedback_dir / "cases.jsonl"))

    return anon_dir


def read_docx_text(path: str | Path) -> str:
    document = Document(path)
    chunks: list[str] = []
    chunks.extend(paragraph.text for paragraph in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    return "\n".join(chunks)


def test_golden_demo_docx_smoke(tmp_path, monkeypatch):
    configure_isolated_runtime(tmp_path, monkeypatch)

    source_docx = build_demo_docx(tmp_path / "golden_synthetic_client_matter.docx")
    project_dictionary = pseudonymize.empty_project_dictionary()

    report, out_path = pseudonymize.process_docx(
        str(source_docx),
        source_docx.name,
        project_dictionary,
        basic_rules(),
    )

    output = read_docx_text(out_path)
    expected_preview = EXPECTED_PREVIEW_PATH.read_text(encoding="utf-8").strip()

    for expected_line in expected_preview.splitlines():
        assert expected_line in output

    for raw_value in [
        "John Carter",
        "04.05.1984",
        "1234 567890",
        "473254765214",
        "Example Street",
        "+7 921 123-45-67",
        "john.carter@example.test",
        "Northwind Analytics",
        "40702810900000000001",
        "4111 1111 1111 1111",
        "legal.team@example.test",
        "+7 495 000-11-22",
    ]:
        assert raw_value not in output

    assert report["summary"]["PERSON"] >= 1
    assert report["summary"]["DATE_BIRTH"] == 1
    assert report["summary"]["PASSPORT"] == 1
    assert report["summary"]["INN"] == 1
    assert report["summary"]["ADDRESS_DETAIL"] == 1
    assert report["summary"]["PHONE"] >= 2
    assert report["summary"]["EMAIL"] >= 2
    assert report["summary"]["ORG_PRIVATE"] == 1
    assert report["summary"]["ACCOUNT"] == 1
    assert report["summary"]["CARD"] == 1

    assert Path(out_path).exists()
