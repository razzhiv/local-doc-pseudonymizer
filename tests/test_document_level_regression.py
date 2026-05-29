from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from PIL import Image

import pseudonymize


def basic_rules() -> dict:
    return {
        "manual_hide": [],
        "manual_allow": [],
        "public_orgs": [],
        "private_org_markers": [],
    }


def configure_isolated_runtime(tmp_path, monkeypatch) -> Path:
    """Redirect generated outputs away from repository working folders."""
    output_dir = tmp_path / "output"
    anon_dir = output_dir / "anonymized"
    reports_dir = output_dir / "reports"
    feedback_dir = tmp_path / "feedback"

    anon_dir.mkdir(parents=True, exist_ok=True)
    reports_dir.mkdir(parents=True, exist_ok=True)
    feedback_dir.mkdir(parents=True, exist_ok=True)

    monkeypatch.setattr(pseudonymize, "OUTPUT_DIR", str(output_dir))
    monkeypatch.setattr(pseudonymize, "ANON_DIR", str(anon_dir))
    monkeypatch.setattr(pseudonymize, "REPORTS_DIR", str(reports_dir))
    monkeypatch.setattr(pseudonymize, "FEEDBACK_DIR", str(feedback_dir))
    monkeypatch.setattr(pseudonymize, "FEEDBACK_CASES_PATH", str(feedback_dir / "cases.jsonl"))

    return anon_dir


def read_docx_text(path: str | Path) -> str:
    document = Document(path)
    chunks: list[str] = []
    chunks.extend(paragraph.text for paragraph in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    return "\n".join(chunks)


def write_text_layer_pdf(path: Path, lines: Iterable[str]) -> None:
    """Create a tiny text-layer PDF without adding a test dependency such as reportlab."""
    commands = ["BT", "/F1 12 Tf", "72 720 Td"]
    first = True
    for line in lines:
        safe = line.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        if first:
            commands.append(f"({safe}) Tj")
            first = False
        else:
            commands.append("0 -18 Td")
            commands.append(f"({safe}) Tj")
    commands.append("ET")
    stream = "\n".join(commands).encode("ascii")

    objects = [
        b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n",
        b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n",
        b"3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>\nendobj\n",
        b"4 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n",
        f"5 0 obj\n<< /Length {len(stream)} >>\nstream\n".encode("ascii")
        + stream
        + b"\nendstream\nendobj\n",
    ]

    data = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = []
    for obj in objects:
        offsets.append(len(data))
        data.extend(obj)

    xref_offset = len(data)
    data.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    data.extend(b"0000000000 65535 f \n")
    for offset in offsets:
        data.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    data.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_offset}\n%%EOF\n".encode("ascii")
    )
    path.write_bytes(bytes(data))


def write_text_plus_blank_pdf(path: Path) -> None:
    """Create a two-page PDF: page 1 has text layer, page 2 has no text."""
    text = "Email: case@example.com".replace("(", "\\(").replace(")", "\\)")
    stream = f"BT\n/F1 12 Tf\n72 720 Td\n({text}) Tj\nET".encode("ascii")

    objects = [
        b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n",
        b"2 0 obj\n<< /Type /Pages /Kids [3 0 R 6 0 R] /Count 2 >>\nendobj\n",
        b"3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>\nendobj\n",
        b"4 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n",
        f"5 0 obj\n<< /Length {len(stream)} >>\nstream\n".encode("ascii")
        + stream
        + b"\nendstream\nendobj\n",
        b"6 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 4 0 R >> >> >>\nendobj\n",
    ]

    data = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = []
    for obj in objects:
        offsets.append(len(data))
        data.extend(obj)

    xref_offset = len(data)
    data.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    data.extend(b"0000000000 65535 f \n")
    for offset in offsets:
        data.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    data.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_offset}\n%%EOF\n".encode("ascii")
    )
    path.write_bytes(bytes(data))


def write_image_only_pdf(path: Path) -> None:
    image = Image.new("RGB", (240, 120), "white")
    image.save(path, "PDF")


def test_docx_paragraph_end_to_end(tmp_path, monkeypatch):
    configure_isolated_runtime(tmp_path, monkeypatch)

    src = tmp_path / "paragraph.docx"
    document = Document()
    document.add_paragraph("ИНН: 473254765214")
    document.add_paragraph("Телефон: +7 921 123-45-67")
    document.save(src)

    project_dictionary = pseudonymize.empty_project_dictionary()
    report, out_path = pseudonymize.process_docx(str(src), src.name, project_dictionary, basic_rules())

    output = read_docx_text(out_path)
    assert "[INN_1]" in output
    assert "[PHONE_1]" in output
    assert "473254765214" not in output
    assert "+7 921 123-45-67" not in output
    assert report["summary"]["INN"] == 1
    assert report["summary"]["PHONE"] == 1


def test_docx_table_context_end_to_end(tmp_path, monkeypatch):
    configure_isolated_runtime(tmp_path, monkeypatch)

    src = tmp_path / "table_context.docx"
    document = Document()
    table = document.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "ФИО"
    table.cell(0, 1).text = "ИНН"
    table.cell(0, 2).text = "Дата рождения"
    table.cell(1, 0).text = "Петров Петр Петрович"
    table.cell(1, 1).text = "473254765214"
    table.cell(1, 2).text = "01.02.1980"
    document.save(src)

    project_dictionary = pseudonymize.empty_project_dictionary()
    report, out_path = pseudonymize.process_docx(str(src), src.name, project_dictionary, basic_rules())

    out = Document(out_path)
    values = [cell.text for cell in out.tables[0].rows[1].cells]
    assert values[0].startswith("[PERSON_")
    assert values[1].startswith("[INN_")
    assert values[2].startswith("[DATE_BIRTH_")
    assert "473254765214" not in "\n".join(values)
    assert "01.02.1980" not in "\n".join(values)
    assert report["summary"]["PERSON"] >= 1
    assert report["summary"]["INN"] == 1
    assert report["summary"]["DATE_BIRTH"] == 1


def test_pdf_text_layer_end_to_end(tmp_path, monkeypatch):
    configure_isolated_runtime(tmp_path, monkeypatch)

    src = tmp_path / "text_layer.pdf"
    write_text_layer_pdf(src, ["Email: user@example.com", "Phone: +7 921 123-45-67"])

    project_dictionary = pseudonymize.empty_project_dictionary()
    report, out_path = pseudonymize.process_pdf(str(src), src.name, project_dictionary, basic_rules())

    assert out_path is not None
    assert Path(out_path).exists()
    assert report["status"] == "processed_text_layer"
    assert report["pdf_extraction"]["pages_total"] == 1
    assert report["pdf_extraction"]["pages_with_text"] == 1
    assert report["pdf_extraction"]["total_extracted_chars"] > 0

    output = read_docx_text(out_path)
    assert "[EMAIL_1]" in output
    assert "[PHONE_1]" in output
    assert "user@example.com" not in output
    assert "+7 921 123-45-67" not in output


def test_pdf_image_only_safe_failure(tmp_path, monkeypatch):
    anon_dir = configure_isolated_runtime(tmp_path, monkeypatch)

    src = tmp_path / "image_only.pdf"
    write_image_only_pdf(src)

    project_dictionary = pseudonymize.empty_project_dictionary()
    report, out_path = pseudonymize.process_pdf(str(src), src.name, project_dictionary, basic_rules())

    assert out_path is None
    assert report["status"] == "not_processed_no_text_layer"
    assert report["pdf_extraction"]["pages_total"] == 1
    assert report["pdf_extraction"]["pages_with_text"] == 0
    assert report["pdf_extraction"]["pages_without_text"] == [1]
    assert report["pdf_extraction"]["total_extracted_chars"] == 0
    assert not list(anon_dir.glob("image_only_anonymized.docx"))


def test_pdf_partial_text_layer_status(tmp_path, monkeypatch):
    configure_isolated_runtime(tmp_path, monkeypatch)

    src = tmp_path / "mixed_text_layer.pdf"
    write_text_plus_blank_pdf(src)

    project_dictionary = pseudonymize.empty_project_dictionary()
    report, out_path = pseudonymize.process_pdf(str(src), src.name, project_dictionary, basic_rules())

    assert out_path is not None
    assert Path(out_path).exists()
    assert report["status"] == "partially_processed_text_layer"
    assert report["pdf_extraction"]["pages_total"] == 2
    assert report["pdf_extraction"]["pages_with_text"] == 1
    assert report["pdf_extraction"]["pages_without_text"] == [2]

    output = read_docx_text(out_path)
    assert "[EMAIL_1]" in output
    assert "case@example.com" not in output
