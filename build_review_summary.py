from pathlib import Path
import json
from datetime import datetime
from collections import Counter, defaultdict

ROOT = Path(__file__).resolve().parent

FEEDBACK_CASES_PATH = ROOT / "feedback" / "cases.jsonl"
REPORT_JSON_PATH = ROOT / "output" / "anonymization_report.json"
REPORTS_DIR = ROOT / "output" / "reports"

LATEST_MD_PATH = REPORTS_DIR / "review_summary_latest.md"


def read_jsonl(path: Path):
    rows = []
    if not path.exists():
        return rows

    with path.open("r", encoding="utf-8") as f:
        for line_no, line in enumerate(f, start=1):
            line = line.strip()
            if not line:
                continue
            try:
                rows.append(json.loads(line))
            except json.JSONDecodeError as e:
                rows.append({
                    "_parse_error": True,
                    "line_no": line_no,
                    "error": str(e),
                    "raw": line,
                })
    return rows


def read_json(path: Path):
    if not path.exists():
        return None
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception as e:
        return {"_parse_error": str(e)}


def safe(value, default=""):
    if value is None:
        return default
    return str(value)


def md_escape_cell(value):
    text = safe(value)
    text = text.replace("|", "\\|")
    text = text.replace("\n", " ")
    return text


def make_markdown(cases, report_json):
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    lines = []
    lines.append("# Review summary")
    lines.append("")
    lines.append(f"_Generated: {now}_")
    lines.append("")
    lines.append("> ВАЖНО: этот отчёт может содержать сведения о структуре документа и чувствительных кейсах. Не отправлять в SaaS и не публиковать без проверки.")
    lines.append("")

    if not cases:
        lines.append("## Статус")
        lines.append("")
        lines.append("В `feedback/cases.jsonl` нет кейсов для review.")
        lines.append("")
        lines.append("Это может означать одно из двух:")
        lines.append("")
        lines.append("1. В последнем прогоне не было `SUSPECT_ENTITY` / `NEEDS_REVIEW`.")
        lines.append("2. `pseudonymize.py` ещё не писал auto-cases или файл был очищен.")
        lines.append("")
        return "\n".join(lines)

    parse_errors = [c for c in cases if c.get("_parse_error")]
    normal_cases = [c for c in cases if not c.get("_parse_error")]

    status_counts = Counter(safe(c.get("status"), "UNKNOWN") for c in normal_cases)
    type_counts = Counter(safe(c.get("entity_type_expected") or c.get("entity_type_detected"), "UNKNOWN") for c in normal_cases)
    reason_counts = Counter(safe(c.get("review_reason"), "UNKNOWN") for c in normal_cases)
    action_counts = Counter(safe(c.get("policy_action"), "UNKNOWN") for c in normal_cases)
    file_counts = Counter(safe(c.get("file_name"), "UNKNOWN") for c in normal_cases)

    lines.append("## Краткий итог")
    lines.append("")
    lines.append(f"- Кейсов для review: **{len(normal_cases)}**")
    if parse_errors:
        lines.append(f"- Ошибок чтения JSONL: **{len(parse_errors)}**")
    lines.append("")

    lines.append("### По статусам")
    lines.append("")
    for k, v in status_counts.most_common():
        lines.append(f"- `{k}`: {v}")
    lines.append("")

    lines.append("### По типам")
    lines.append("")
    for k, v in type_counts.most_common():
        lines.append(f"- `{k}`: {v}")
    lines.append("")

    lines.append("### По действиям политики")
    lines.append("")
    for k, v in action_counts.most_common():
        lines.append(f"- `{k}`: {v}")
    lines.append("")

    lines.append("### По причинам review")
    lines.append("")
    for k, v in reason_counts.most_common():
        lines.append(f"- `{k}`: {v}")
    lines.append("")

    lines.append("### По файлам")
    lines.append("")
    for k, v in file_counts.most_common():
        lines.append(f"- `{k}`: {v}")
    lines.append("")

    lines.append("## Кейсы для решения человеком")
    lines.append("")
    lines.append("| № | case_id | file | block | status | type | action | reason | safe_context | recommended |")
    lines.append("|---:|---|---|---|---|---|---|---|---|---|")

    for i, c in enumerate(normal_cases, start=1):
        type_value = c.get("entity_type_expected") or c.get("entity_type_detected") or ""
        lines.append(
            "| "
            + " | ".join([
                str(i),
                md_escape_cell(c.get("case_id")),
                md_escape_cell(c.get("file_name")),
                md_escape_cell(c.get("block")),
                md_escape_cell(c.get("status")),
                md_escape_cell(type_value),
                md_escape_cell(c.get("policy_action")),
                md_escape_cell(c.get("review_reason")),
                md_escape_cell(c.get("safe_context")),
                md_escape_cell(c.get("recommended_decision")),
            ])
            + " |"
        )

    lines.append("")
    lines.append("## Что сделать дальше")
    lines.append("")
    lines.append("1. Открыть или создать Excel-review:")
    lines.append("")
    lines.append("```powershell")
    lines.append("python review_tool.py export")
    lines.append("```")
    lines.append("")
    lines.append("2. Проверить решения на листе `Решения`.")
    lines.append("")
    lines.append("3. Сохранить Excel и импортировать решения:")
    lines.append("")
    lines.append("```powershell")
    lines.append("python review_tool.py import")
    lines.append("```")
    lines.append("")
    lines.append("4. После появления candidate-rules / regression-cases запускать regression:")
    lines.append("")
    lines.append("```powershell")
    lines.append("python run_regression_tests.py run")
    lines.append("```")
    lines.append("")

    if report_json:
        lines.append("## Связь с anonymization_report.json")
        lines.append("")
        if isinstance(report_json, dict) and report_json.get("_parse_error"):
            lines.append(f"`output/anonymization_report.json` найден, но не прочитан: `{report_json['_parse_error']}`")
        else:
            keys = ", ".join(sorted(report_json.keys())) if isinstance(report_json, dict) else type(report_json).__name__
            lines.append("`output/anonymization_report.json` найден.")
            lines.append("")
            lines.append(f"Верхнеуровневые ключи: `{keys}`")
        lines.append("")

    if parse_errors:
        lines.append("## Ошибки чтения JSONL")
        lines.append("")
        for e in parse_errors:
            lines.append(f"- line {e.get('line_no')}: {e.get('error')}")
        lines.append("")

    return "\n".join(lines)


def write_optional_docx(md_text: str, docx_path: Path):
    try:
        from docx import Document
    except ImportError:
        return False, "python-docx не установлен"

    doc = Document()
    for line in md_text.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.startswith("> "):
            p = doc.add_paragraph(line[2:].strip())
            p.style = doc.styles["Intense Quote"] if "Intense Quote" in doc.styles else doc.styles["Normal"]
        elif line.startswith("|"):
            # Таблицы в DOCX на этом этапе не парсим, оставляем как текст.
            doc.add_paragraph(line)
        elif line.startswith("```"):
            continue
        else:
            doc.add_paragraph(line)

    doc.save(docx_path)
    return True, "ok"


def main():
    print("=== BUILD REVIEW SUMMARY ===")

    REPORTS_DIR.mkdir(parents=True, exist_ok=True)

    cases = read_jsonl(FEEDBACK_CASES_PATH)
    report_json = read_json(REPORT_JSON_PATH)

    md = make_markdown(cases, report_json)

    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    md_path = REPORTS_DIR / f"review_summary_{timestamp}.md"
    docx_path = REPORTS_DIR / f"review_summary_{timestamp}.docx"

    md_path.write_text(md, encoding="utf-8")
    LATEST_MD_PATH.write_text(md, encoding="utf-8")

    print(f"MD отчёт: {md_path}")
    print(f"Latest MD: {LATEST_MD_PATH}")

    ok, message = write_optional_docx(md, docx_path)
    if ok:
        print(f"DOCX отчёт: {docx_path}")
    else:
        print(f"DOCX не создан: {message}")

    print("")
    print("Кратко:")
    print(f"- Кейсов в feedback/cases.jsonl: {len([c for c in cases if not c.get('_parse_error')])}")
    print("")
    print("Следующий шаг:")
    print("python review_tool.py export")


if __name__ == "__main__":
    main()
